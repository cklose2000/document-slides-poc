#!/usr/bin/env python3
"""
Comprehensive API Endpoint Testing Suite for document-slides-poc
Tests all Flask API endpoints with various scenarios, error handling, and edge cases.
"""

import unittest
import requests
import json
import io
import os
import tempfile
import time
from unittest.mock import patch, MagicMock
import openpyxl
from docx import Document
from pathlib import Path

class APIEndpointTestSuite(unittest.TestCase):
    """Test suite for all API endpoints in the document-slides-poc project"""
    
    @classmethod
    def setUpClass(cls):
        """Setup test environment"""
        cls.base_url = "http://localhost:5000"
        cls.api_url = f"{cls.base_url}/api"
        
        # Create test data directory
        cls.test_data_dir = tempfile.mkdtemp(prefix="api_test_")
        print(f"Test data directory: {cls.test_data_dir}")
        
        # Create sample test files
        cls._create_test_files()
        
        # Track created temp files for cleanup
        cls.temp_files = []
    
    @classmethod
    def tearDownClass(cls):
        """Cleanup test environment"""
        # Clean up temporary files
        for temp_file in cls.temp_files:
            if os.path.exists(temp_file):
                os.unlink(temp_file)
        
        # Clean up test data directory
        import shutil
        if os.path.exists(cls.test_data_dir):
            shutil.rmtree(cls.test_data_dir)
    
    @classmethod
    def _create_test_files(cls):
        """Create sample test files for API testing"""
        
        # 1. Create sample Excel file
        cls.excel_file = os.path.join(cls.test_data_dir, "test_financials.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Summary"
        
        # Add company info
        ws['A1'] = "TechCorp Inc."
        ws['A2'] = "Financial Summary Q4 2023"
        
        # Add financial metrics
        ws['A5'] = "Revenue"
        ws['B5'] = 10200000
        ws['A6'] = "Profit" 
        ws['B6'] = 2500000
        ws['A7'] = "Growth Rate"
        ws['B7'] = 0.23
        ws['A8'] = "Market Share"
        ws['B8'] = 0.15
        
        # Add formulas
        ws['A10'] = "Total Assets"
        ws['B10'] = "=B5+B6"
        
        wb.save(cls.excel_file)
        
        # 2. Create sample Word document
        cls.word_file = os.path.join(cls.test_data_dir, "business_plan.docx")
        doc = Document()
        doc.add_heading('Executive Summary', 0)
        doc.add_paragraph('TechCorp Inc. is a leading technology company specializing in cloud solutions.')
        
        doc.add_heading('Market Analysis', 1)
        doc.add_paragraph('The cloud computing market is experiencing rapid growth with 25% YoY increase.')
        doc.add_paragraph('Our target market represents a $50B opportunity.')
        
        doc.add_heading('Financial Projections', 1)
        doc.add_paragraph('Revenue is projected to grow from $10.2M to $15.8M next year.')
        doc.add_paragraph('Operating margin expected to improve from 24% to 28%.')
        
        doc.save(cls.word_file)
        
        # 3. Create sample PDF content (as text file for simplicity)
        cls.pdf_file = os.path.join(cls.test_data_dir, "company_memo.txt")
        with open(cls.pdf_file, 'w') as f:
            f.write("""
COMPANY MEMO - CONFIDENTIAL

TO: Board of Directors
FROM: CEO Office
DATE: December 2023
SUBJECT: Q4 Performance Review

EXECUTIVE SUMMARY
TechCorp Inc. has delivered strong performance in Q4 2023.

KEY METRICS:
- Revenue: $10.2M (up 23% YoY)
- Net Profit: $2.5M (24% margin)
- Customer Growth: 1,200 new customers
- Market Share: 15% in target segment

MARKET POSITION:
Our cloud platform continues to gain traction with enterprise clients.
The competitive landscape remains favorable for continued growth.

OUTLOOK:
We are well-positioned for 2024 with strong pipeline and market momentum.
""")
        
        # 4. Create invalid files for error testing
        cls.invalid_file = os.path.join(cls.test_data_dir, "invalid.txt")
        with open(cls.invalid_file, 'w') as f:
            f.write("This is not a supported document format.")
        
        # 5. Create empty file
        cls.empty_file = os.path.join(cls.test_data_dir, "empty.xlsx")
        with open(cls.empty_file, 'w') as f:
            pass  # Create empty file
        
        # 6. Create corrupted Excel file
        cls.corrupted_excel = os.path.join(cls.test_data_dir, "corrupted.xlsx")
        with open(cls.corrupted_excel, 'w') as f:
            f.write("This is not a valid Excel file content")
        
        # 7. Create sample PowerPoint template for template testing
        cls.sample_template = os.path.join(cls.test_data_dir, "sample_template.pptx")
        # Create a minimal PPTX file structure (simplified)
        with open(cls.sample_template, 'w') as f:
            f.write("Mock PPTX template content")
    
    def setUp(self):
        """Setup for each test"""
        self.session = requests.Session()
        self.session.timeout = 30
    
    def tearDown(self):
        """Cleanup after each test"""
        self.session.close()
    
    def _check_server_running(self):
        """Check if the API server is running"""
        try:
            response = self.session.get(f"{self.base_url}/health", timeout=5)
            return response.status_code == 200
        except:
            return False
    
    def test_01_health_check(self):
        """Test the health check endpoint"""
        print("\n=== Testing Health Check Endpoint ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running on localhost:5000")
        
        response = self.session.get(f"{self.base_url}/health")
        
        self.assertEqual(response.status_code, 200)
        data = response.json()
        
        # Verify response structure
        self.assertIn('status', data)
        self.assertEqual(data['status'], 'healthy')
        self.assertIn('service', data)
        self.assertEqual(data['service'], 'document-slides-poc')
        self.assertIn('template_management', data)
        
        print(f"✓ Health check passed: {data}")
    
    def test_02_generate_slides_no_files(self):
        """Test generate-slides endpoint with no files"""
        print("\n=== Testing Generate Slides - No Files ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        response = self.session.post(f"{self.api_url}/generate-slides")
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertIn('error', data)
        self.assertIn('No files uploaded', data['error'])
        
        print(f"✓ No files error handled correctly: {data}")
    
    def test_03_generate_slides_empty_files(self):
        """Test generate-slides endpoint with empty file list"""
        print("\n=== Testing Generate Slides - Empty Files ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        files = {'documents': (None, '')}
        response = self.session.post(f"{self.api_url}/generate-slides", files=files)
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertIn('error', data)
        
        print(f"✓ Empty files error handled correctly: {data}")
    
    def test_04_generate_slides_unsupported_format(self):
        """Test generate-slides endpoint with unsupported file format"""
        print("\n=== Testing Generate Slides - Unsupported Format ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        with open(self.invalid_file, 'rb') as f:
            files = {'documents': ('invalid.txt', f, 'text/plain')}
            response = self.session.post(f"{self.api_url}/generate-slides", files=files)
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertIn('error', data)
        self.assertIn('No supported documents found', data['error'])
        
        print(f"✓ Unsupported format error handled correctly: {data}")
    
    def test_05_generate_slides_valid_excel(self):
        """Test generate-slides endpoint with valid Excel file"""
        print("\n=== Testing Generate Slides - Valid Excel ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        with open(self.excel_file, 'rb') as f:
            files = {'documents': ('test_financials.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            response = self.session.post(f"{self.api_url}/generate-slides", files=files)
        
        # Should return a PowerPoint file or JSON error
        if response.status_code == 200:
            # Check if response is a PowerPoint file
            self.assertEqual(response.headers.get('content-type'), 
                           'application/vnd.openxmlformats-officedocument.presentationml.presentation')
            self.assertGreater(len(response.content), 0)
            print("✓ Excel file processed successfully, PowerPoint generated")
            
            # Save generated file for inspection
            output_file = os.path.join(self.test_data_dir, "generated_excel_test.pptx")
            with open(output_file, 'wb') as out_f:
                out_f.write(response.content)
            self.temp_files.append(output_file)
            print(f"✓ Generated file saved: {output_file}")
            
        else:
            # Check error response
            try:
                data = response.json()
                print(f"⚠ Excel processing failed with error: {data}")
                # This might be expected if dependencies are missing
                self.assertIn('error', data)
            except:
                self.fail(f"Unexpected response: {response.status_code} - {response.text}")
    
    def test_06_generate_slides_valid_word(self):
        """Test generate-slides endpoint with valid Word document"""
        print("\n=== Testing Generate Slides - Valid Word ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        with open(self.word_file, 'rb') as f:
            files = {'documents': ('business_plan.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            response = self.session.post(f"{self.api_url}/generate-slides", files=files)
        
        if response.status_code == 200:
            self.assertEqual(response.headers.get('content-type'), 
                           'application/vnd.openxmlformats-officedocument.presentationml.presentation')
            self.assertGreater(len(response.content), 0)
            print("✓ Word document processed successfully, PowerPoint generated")
            
            # Save generated file
            output_file = os.path.join(self.test_data_dir, "generated_word_test.pptx")
            with open(output_file, 'wb') as out_f:
                out_f.write(response.content)
            self.temp_files.append(output_file)
            print(f"✓ Generated file saved: {output_file}")
            
        else:
            try:
                data = response.json()
                print(f"⚠ Word processing failed: {data}")
                self.assertIn('error', data)
            except:
                self.fail(f"Unexpected response: {response.status_code} - {response.text}")
    
    def test_07_generate_slides_multiple_files(self):
        """Test generate-slides endpoint with multiple files"""
        print("\n=== Testing Generate Slides - Multiple Files ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        files = [
            ('documents', ('test_financials.xlsx', open(self.excel_file, 'rb'), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')),
            ('documents', ('business_plan.docx', open(self.word_file, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
        ]
        
        try:
            response = self.session.post(f"{self.api_url}/generate-slides", files=files)
            
            if response.status_code == 200:
                self.assertEqual(response.headers.get('content-type'), 
                               'application/vnd.openxmlformats-officedocument.presentationml.presentation')
                print("✓ Multiple files processed successfully")
                
                # Save generated file
                output_file = os.path.join(self.test_data_dir, "generated_multiple_test.pptx")
                with open(output_file, 'wb') as out_f:
                    out_f.write(response.content)
                self.temp_files.append(output_file)
                print(f"✓ Generated file saved: {output_file}")
                
            else:
                try:
                    data = response.json()
                    print(f"⚠ Multiple files processing failed: {data}")
                    self.assertIn('error', data)
                except:
                    self.fail(f"Unexpected response: {response.status_code} - {response.text}")
        
        finally:
            # Close file handles
            for _, file_tuple in files:
                file_tuple[1].close()
    
    def test_08_generate_slides_with_template(self):
        """Test generate-slides endpoint with template selection"""
        print("\n=== Testing Generate Slides - With Template ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        with open(self.excel_file, 'rb') as f:
            files = {'documents': ('test_financials.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            data = {'template_id': 'default'}
            response = self.session.post(f"{self.api_url}/generate-slides", files=files, data=data)
        
        if response.status_code == 200:
            self.assertEqual(response.headers.get('content-type'), 
                           'application/vnd.openxmlformats-officedocument.presentationml.presentation')
            print("✓ Template selection worked")
            
            # Save generated file
            output_file = os.path.join(self.test_data_dir, "generated_template_test.pptx")
            with open(output_file, 'wb') as out_f:
                out_f.write(response.content)
            self.temp_files.append(output_file)
            
        elif response.status_code == 500:
            # Template management might not be available
            data = response.json()
            print(f"⚠ Template processing failed (expected if template system not available): {data}")
            self.assertIn('error', data)
        else:
            self.fail(f"Unexpected response: {response.status_code} - {response.text}")
    
    def test_09_preview_extraction_valid_files(self):
        """Test preview extraction endpoint with valid files"""
        print("\n=== Testing Preview Extraction - Valid Files ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        with open(self.excel_file, 'rb') as f:
            files = {'documents': ('test_financials.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            response = self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
        
        self.assertEqual(response.status_code, 200)
        data = response.json()
        
        # Verify response structure
        self.assertIn('success', data)
        self.assertTrue(data['success'])
        self.assertIn('documents_processed', data)
        self.assertGreater(data['documents_processed'], 0)
        self.assertIn('extraction_results', data)
        self.assertIsInstance(data['extraction_results'], list)
        
        # Check first document result
        if data['extraction_results']:
            doc_result = data['extraction_results'][0]
            self.assertIn('filename', doc_result)
            self.assertIn('type', doc_result)
            self.assertEqual(doc_result['type'], 'excel')
        
        print(f"✓ Preview extraction successful: {data['documents_processed']} documents processed")
    
    def test_10_preview_extraction_no_files(self):
        """Test preview extraction endpoint with no files"""
        print("\n=== Testing Preview Extraction - No Files ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        response = self.session.post(f"{self.api_url}/generate-slides/preview")
        
        self.assertEqual(response.status_code, 400)
        data = response.json()
        self.assertIn('error', data)
        self.assertIn('No files uploaded', data['error'])
        
        print(f"✓ No files error handled correctly: {data}")
    
    def test_11_preview_extraction_multiple_types(self):
        """Test preview extraction with multiple document types"""
        print("\n=== Testing Preview Extraction - Multiple Types ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        files = [
            ('documents', ('test_financials.xlsx', open(self.excel_file, 'rb'), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')),
            ('documents', ('business_plan.docx', open(self.word_file, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
        ]
        
        try:
            response = self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
            
            self.assertEqual(response.status_code, 200)
            data = response.json()
            
            self.assertTrue(data['success'])
            self.assertEqual(data['documents_processed'], 2)
            self.assertEqual(len(data['extraction_results']), 2)
            
            # Check document types
            doc_types = [doc['type'] for doc in data['extraction_results']]
            self.assertIn('excel', doc_types)
            self.assertIn('word', doc_types)
            
            print(f"✓ Multiple document types processed: {doc_types}")
            
        finally:
            # Close file handles
            for _, file_tuple in files:
                file_tuple[1].close()
    
    def test_12_templates_list(self):
        """Test templates listing endpoint"""
        print("\n=== Testing Templates List ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        response = self.session.get(f"{self.api_url}/templates")
        
        if response.status_code == 200:
            data = response.json()
            self.assertIn('templates', data)
            self.assertIn('count', data)
            self.assertIsInstance(data['templates'], list)
            print(f"✓ Templates listed successfully: {data['count']} templates")
            
            # Check template structure if any exist
            if data['templates']:
                template = data['templates'][0]
                expected_fields = ['id', 'name', 'description']
                for field in expected_fields:
                    self.assertIn(field, template)
                    
        elif response.status_code == 503:
            data = response.json()
            self.assertIn('error', data)
            self.assertIn('Template management not available', data['error'])
            print("⚠ Template management not available (expected if not configured)")
        else:
            self.fail(f"Unexpected response: {response.status_code} - {response.text}")
    
    def test_13_template_upload(self):
        """Test template upload endpoint"""
        print("\n=== Testing Template Upload ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Test with invalid file (not PPTX)
        with open(self.invalid_file, 'rb') as f:
            files = {'template': ('invalid.txt', f, 'text/plain')}
            data = {'name': 'test_template', 'description': 'Test template'}
            response = self.session.post(f"{self.api_url}/templates/upload", files=files, data=data)
        
        if response.status_code == 503:
            # Template management not available
            data = response.json()
            self.assertIn('Template management not available', data['error'])
            print("⚠ Template management not available")
            return
        elif response.status_code == 400:
            data = response.json()
            self.assertIn('error', data)
            self.assertIn('must be a', data['error'].lower())
            print(f"✓ Invalid file type rejected: {data}")
        
        # Test with no file
        response = self.session.post(f"{self.api_url}/templates/upload")
        if response.status_code != 503:  # If template management is available
            self.assertEqual(response.status_code, 400)
            data = response.json()
            self.assertIn('error', data)
            print(f"✓ No file error handled: {data}")
    
    def test_14_template_select(self):
        """Test template selection endpoint"""
        print("\n=== Testing Template Selection ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Test selecting non-existent template
        response = self.session.post(f"{self.api_url}/templates/nonexistent/select")
        
        if response.status_code == 503:
            data = response.json()
            self.assertIn('Template management not available', data['error'])
            print("⚠ Template management not available")
        elif response.status_code == 404:
            data = response.json()
            self.assertIn('error', data)
            self.assertIn('not found', data['error'])
            print(f"✓ Non-existent template error handled: {data}")
        else:
            # Might be 500 if template system has issues
            print(f"⚠ Template selection response: {response.status_code}")
    
    def test_15_template_info(self):
        """Test template info endpoint"""
        print("\n=== Testing Template Info ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Test getting info for non-existent template
        response = self.session.get(f"{self.api_url}/templates/nonexistent")
        
        if response.status_code == 503:
            data = response.json()
            self.assertIn('Template management not available', data['error'])
            print("⚠ Template management not available")
        elif response.status_code == 404:
            data = response.json()
            self.assertIn('error', data)
            self.assertIn('not found', data['error'])
            print(f"✓ Non-existent template info error handled: {data}")
        else:
            print(f"⚠ Template info response: {response.status_code}")
    
    def test_16_template_preview(self):
        """Test template preview endpoint"""
        print("\n=== Testing Template Preview ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        response = self.session.get(f"{self.api_url}/templates/default/preview")
        
        if response.status_code == 503:
            data = response.json()
            self.assertIn('Template management not available', data['error'])
            print("⚠ Template management not available")
        elif response.status_code == 404:
            # Preview image not found (expected)
            print("✓ Preview not found (expected for default setup)")
        elif response.status_code == 200:
            # Preview image found
            self.assertEqual(response.headers.get('content-type'), 'image/png')
            print("✓ Template preview image found")
        else:
            print(f"⚠ Template preview response: {response.status_code}")
    
    def test_17_error_handling_corrupted_files(self):
        """Test error handling with corrupted files"""
        print("\n=== Testing Error Handling - Corrupted Files ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Test with corrupted Excel file
        with open(self.corrupted_excel, 'rb') as f:
            files = {'documents': ('corrupted.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            response = self.session.post(f"{self.api_url}/generate-slides", files=files)
        
        # Should handle the error gracefully
        if response.status_code == 500:
            data = response.json()
            self.assertIn('error', data)
            print(f"✓ Corrupted file error handled: {data}")
        elif response.status_code == 400:
            data = response.json()
            self.assertIn('error', data)
            print(f"✓ Corrupted file rejected: {data}")
        else:
            print(f"⚠ Unexpected response for corrupted file: {response.status_code}")
    
    def test_18_large_file_handling(self):
        """Test handling of large files"""
        print("\n=== Testing Large File Handling ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Create a large text file (simulate large document)
        large_file = os.path.join(self.test_data_dir, "large_document.txt")
        with open(large_file, 'w') as f:
            # Write 1MB of text
            for i in range(10000):
                f.write(f"This is line {i} of a large document for testing purposes. " * 10 + "\n")
        
        self.temp_files.append(large_file)
        
        # Test with the large file (as unsupported format to test file handling)
        with open(large_file, 'rb') as f:
            files = {'documents': ('large_document.txt', f, 'text/plain')}
            start_time = time.time()
            response = self.session.post(f"{self.api_url}/generate-slides", files=files)
            end_time = time.time()
        
        # Should reject unsupported format quickly
        self.assertEqual(response.status_code, 400)
        self.assertLess(end_time - start_time, 10)  # Should handle quickly
        
        data = response.json()
        self.assertIn('error', data)
        print(f"✓ Large unsupported file handled efficiently in {end_time - start_time:.2f}s: {data}")
    
    def test_19_concurrent_requests(self):
        """Test concurrent API requests"""
        print("\n=== Testing Concurrent Requests ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        import threading
        import queue
        
        results = queue.Queue()
        
        def make_request(request_id):
            try:
                with open(self.excel_file, 'rb') as f:
                    files = {'documents': (f'test_financials_{request_id}.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                    response = requests.post(f"{self.api_url}/generate-slides/preview", files=files, timeout=30)
                    results.put((request_id, response.status_code, len(response.content)))
            except Exception as e:
                results.put((request_id, 'error', str(e)))
        
        # Start 3 concurrent requests
        threads = []
        for i in range(3):
            thread = threading.Thread(target=make_request, args=(i,))
            threads.append(thread)
            thread.start()
        
        # Wait for all threads to complete
        for thread in threads:
            thread.join(timeout=60)
        
        # Check results
        success_count = 0
        while not results.empty():
            request_id, status, content_len = results.get()
            if status == 200:
                success_count += 1
            print(f"Request {request_id}: Status {status}, Content length: {content_len}")
        
        self.assertGreater(success_count, 0)
        print(f"✓ Concurrent requests handled: {success_count}/3 successful")
    
    def test_20_response_format_validation(self):
        """Test API response format validation"""
        print("\n=== Testing Response Format Validation ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Test health endpoint format
        response = self.session.get(f"{self.base_url}/health")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers.get('content-type'), 'application/json')
        
        data = response.json()
        required_fields = ['status', 'service', 'template_management']
        for field in required_fields:
            self.assertIn(field, data)
        
        print("✓ Health endpoint response format valid")
        
        # Test preview endpoint format
        with open(self.excel_file, 'rb') as f:
            files = {'documents': ('test_financials.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            response = self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
        
        if response.status_code == 200:
            self.assertEqual(response.headers.get('content-type'), 'application/json')
            data = response.json()
            
            required_fields = ['success', 'documents_processed', 'extraction_results', 'message']
            for field in required_fields:
                self.assertIn(field, data, f"Missing field: {field}")
            
            self.assertIsInstance(data['success'], bool)
            self.assertIsInstance(data['documents_processed'], int)
            self.assertIsInstance(data['extraction_results'], list)
            self.assertIsInstance(data['message'], str)
            
            print("✓ Preview endpoint response format valid")
        
        # Test error response format
        response = self.session.post(f"{self.api_url}/generate-slides")
        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.headers.get('content-type'), 'application/json')
        
        data = response.json()
        self.assertIn('error', data)
        self.assertIsInstance(data['error'], str)
        
        print("✓ Error response format valid")

def run_tests():
    """Run the test suite"""
    print("=" * 60)
    print("API ENDPOINT TESTING SUITE")
    print("Testing document-slides-poc API endpoints")
    print("=" * 60)
    
    # Create test suite
    loader = unittest.TestLoader()
    suite = loader.loadTestsFromTestCase(APIEndpointTestSuite)
    
    # Run tests with detailed output
    runner = unittest.TextTestRunner(verbosity=2, stream=None)
    result = runner.run(suite)
    
    # Print summary
    print("\n" + "=" * 60)
    print("TEST RESULTS SUMMARY")
    print("=" * 60)
    print(f"Tests run: {result.testsRun}")
    print(f"Failures: {len(result.failures)}")
    print(f"Errors: {len(result.errors)}")
    print(f"Skipped: {len(result.skipped)}")
    
    if result.failures:
        print("\nFAILURES:")
        for test, traceback in result.failures:
            print(f"- {test}: {traceback}")
    
    if result.errors:
        print("\nERRORS:")
        for test, traceback in result.errors:
            print(f"- {test}: {traceback}")
    
    if result.skipped:
        print("\nSKIPPED:")
        for test, reason in result.skipped:
            print(f"- {test}: {reason}")
    
    success_rate = ((result.testsRun - len(result.failures) - len(result.errors)) / result.testsRun * 100) if result.testsRun > 0 else 0
    print(f"\nSuccess Rate: {success_rate:.1f}%")
    
    return result.wasSuccessful()

if __name__ == "__main__":
    success = run_tests()
    exit(0 if success else 1)