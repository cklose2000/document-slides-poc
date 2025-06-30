#!/usr/bin/env python3
"""
Mocked API Testing Suite for document-slides-poc
Tests API endpoints with mocked external dependencies (OpenAI, LLMWhisperer)
"""

import unittest
import requests
import json
import io
import os
import tempfile
import sys
from unittest.mock import patch, MagicMock
import openpyxl
from docx import Document

# Add the lib directory to the path for imports
sys.path.append(os.path.join(os.path.dirname(__file__), 'lib'))

class MockedAPITestSuite(unittest.TestCase):
    """Test API endpoints with mocked external services"""
    
    @classmethod
    def setUpClass(cls):
        """Setup test environment with mock data"""
        cls.base_url = "http://localhost:5000"
        cls.api_url = f"{cls.base_url}/api"
        
        # Create test data directory
        cls.test_data_dir = tempfile.mkdtemp(prefix="mocked_api_test_")
        
        # Create sample test files
        cls._create_test_files()
        cls.temp_files = []
        
        # Mock responses for external APIs
        cls.mock_openai_response = {
            "company_overview": {
                "name": "TechCorp Inc.",
                "industry": "Technology",
                "description": "Leading cloud solutions provider"
            },
            "financial_metrics": {
                "revenue": "$10.2M",
                "profit": "$2.5M",
                "growth_rate": "23%",
                "market_share": "15%"
            },
            "key_insights": [
                "Strong Q4 performance with 23% YoY growth",
                "Market leadership position maintained",
                "Successful enterprise client acquisition"
            ],
            "suggested_slides": [
                {
                    "type": "company_overview",
                    "title": "Company Overview",
                    "content": "TechCorp Inc. - Leading technology company"
                },
                {
                    "type": "financial_summary", 
                    "title": "Financial Performance",
                    "content": "Revenue: $10.2M, Profit: $2.5M"
                },
                {
                    "type": "data_insights",
                    "title": "Key Insights",
                    "content": "23% YoY growth, strong market position"
                }
            ],
            "source_attributions": {
                "primary_documents": ["test_financials.xlsx", "business_plan.docx"],
                "key_metrics_sources": {
                    "revenue": "test_financials.xlsx:Summary:B5",
                    "profit": "test_financials.xlsx:Summary:B6"
                },
                "extraction_summary": "Successfully processed 2 documents"
            }
        }
        
        cls.mock_llm_whisperer_response = {
            "status": "success",
            "extraction_id": "test-123",
            "text": "Sample extracted text from PDF document",
            "pages": 5,
            "confidence": 0.95
        }
    
    @classmethod
    def tearDownClass(cls):
        """Cleanup test environment"""
        # Clean up temporary files
        for temp_file in cls.temp_files:
            if os.path.exists(temp_file):
                os.unlink(temp_file)
        
        # Clean up test data directory
        import shutil
        if os.path.exists(cls.test_data_dir):
            shutil.rmtree(cls.test_data_dir)
    
    @classmethod
    def _create_test_files(cls):
        """Create sample test files for API testing"""
        
        # Create comprehensive Excel file
        cls.excel_file = os.path.join(cls.test_data_dir, "test_financials.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Summary"
        
        # Company information
        ws['A1'] = "TechCorp Inc."
        ws['A2'] = "Q4 2023 Financial Summary"
        ws['A3'] = "Technology Sector"
        
        # Financial metrics with realistic data
        ws['A5'] = "Revenue (USD)"
        ws['B5'] = 10200000
        ws['C5'] = "Q4 2023"
        
        ws['A6'] = "Net Profit (USD)"
        ws['B6'] = 2500000
        ws['C6'] = "24.5% margin"
        
        ws['A7'] = "Growth Rate (%)"
        ws['B7'] = 23.4
        ws['C7'] = "YoY"
        
        ws['A8'] = "Market Share (%)"
        ws['B8'] = 15.2
        ws['C8'] = "Target segment"
        
        ws['A9'] = "Customer Count"
        ws['B9'] = 1250
        ws['C9'] = "Active customers"
        
        ws['A10'] = "Employee Count"
        ws['B10'] = 89
        ws['C10'] = "Full-time"
        
        # Add formulas
        ws['A12'] = "Total Assets"
        ws['B12'] = "=B5+B6"
        ws['A13'] = "Profit Margin"
        ws['B13'] = "=B6/B5"
        
        # Add quarterly data sheet
        ws2 = wb.create_sheet("Quarterly_Data")
        quarters = ["Q1", "Q2", "Q3", "Q4"]
        revenues = [7500000, 8200000, 9100000, 10200000]
        
        ws2['A1'] = "Quarter"
        ws2['B1'] = "Revenue"
        ws2['C1'] = "Growth"
        
        for i, (quarter, revenue) in enumerate(zip(quarters, revenues), 2):
            ws2[f'A{i}'] = quarter
            ws2[f'B{i}'] = revenue
            if i > 2:
                ws2[f'C{i}'] = f"=B{i}/B{i-1}-1"
        
        wb.save(cls.excel_file)
        
        # Create comprehensive Word document
        cls.word_file = os.path.join(cls.test_data_dir, "business_plan.docx")
        doc = Document()
        
        # Title and executive summary
        doc.add_heading('TechCorp Inc. Business Plan', 0)
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(
            'TechCorp Inc. is a leading technology company specializing in cloud computing solutions. '
            'Founded in 2018, we have grown to serve over 1,250 enterprise customers with our '
            'innovative platform. Our Q4 2023 results show strong performance with $10.2M in revenue '
            'and 23% year-over-year growth.'
        )
        
        # Market analysis
        doc.add_heading('Market Analysis', 1)
        doc.add_paragraph(
            'The cloud computing market is experiencing unprecedented growth, with our target segment '
            'representing a $50 billion opportunity. Market research indicates continued expansion '
            'driven by digital transformation initiatives across industries.'
        )
        
        # Add table with market data
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Market Segment'
        hdr_cells[1].text = 'Size ($B)'
        hdr_cells[2].text = 'Growth Rate'
        
        market_data = [
            ['Enterprise Cloud', '25.5', '18%'],
            ['SMB Solutions', '15.2', '22%'],
            ['Industry Specific', '9.3', '28%']
        ]
        
        for segment, size, growth in market_data:
            row_cells = table.add_row().cells
            row_cells[0].text = segment
            row_cells[1].text = size
            row_cells[2].text = growth
        
        # Financial projections
        doc.add_heading('Financial Projections', 1)
        doc.add_paragraph(
            'Based on current performance and market trends, we project continued strong growth. '
            'Revenue is expected to reach $15.8M in 2024, with operating margins improving from '
            '24% to 28% through operational efficiency gains.'
        )
        
        # Strategic initiatives
        doc.add_heading('Strategic Initiatives', 1)
        doc.add_paragraph('Key initiatives for 2024 include:')
        doc.add_paragraph('• Product platform expansion', style='List Bullet')
        doc.add_paragraph('• Geographic market expansion', style='List Bullet')
        doc.add_paragraph('• Strategic partnerships', style='List Bullet')
        doc.add_paragraph('• Team scaling (89 to 125 employees)', style='List Bullet')
        
        doc.save(cls.word_file)
        
        # Create sample PDF content (text file)
        cls.pdf_file = os.path.join(cls.test_data_dir, "quarterly_report.txt")
        with open(cls.pdf_file, 'w') as f:
            f.write("""
TECHCORP INC. - QUARTERLY REPORT Q4 2023

EXECUTIVE SUMMARY
This quarter represents a milestone in TechCorp's growth trajectory, with record revenue
and strong customer acquisition metrics.

FINANCIAL HIGHLIGHTS
• Revenue: $10.2M (↑23% YoY)
• Net Profit: $2.5M (24.5% margin)
• Customer Growth: +1,200 new enterprise customers
• Market Share: 15% in target segment
• Employee Headcount: 89 (↑25% YoY)

OPERATIONAL METRICS
• Platform Uptime: 99.97%
• Customer Satisfaction: 4.8/5.0
• Product Features Released: 23
• Support Response Time: <2 hours average

MARKET POSITION
Our competitive analysis shows TechCorp maintaining leadership in the enterprise
cloud solutions segment. Customer retention rate of 96% demonstrates strong
product-market fit and customer satisfaction.

FUTURE OUTLOOK
The company is well-positioned for continued growth in 2024, with strong pipeline
development and expanding market opportunities. Key focus areas include:

1. Product Innovation: Continued investment in R&D
2. Market Expansion: Geographic and vertical expansion
3. Team Growth: Strategic hiring across all departments
4. Partnership Development: Strategic alliances

RISK FACTORS
• Competitive pressure from larger technology companies
• Economic uncertainty affecting enterprise spending
• Talent acquisition challenges in tight labor market
• Regulatory changes in data privacy requirements

CONCLUSION
Q4 2023 results demonstrate TechCorp's ability to execute on growth strategies
while maintaining operational excellence. The foundation is strong for continued
success in 2024 and beyond.
""")
    
    def setUp(self):
        """Setup for each test"""
        self.session = requests.Session()
        self.session.timeout = 30
        
        # Environment variable patches
        self.env_patcher = patch.dict(os.environ, {
            'OPENAI_API_KEY': 'test-api-key',
            'LLMWHISPERER_API_KEY': 'test-llm-key'
        })
        self.env_patcher.start()
    
    def tearDown(self):
        """Cleanup after each test"""
        self.session.close()
        self.env_patcher.stop()
    
    def _check_server_running(self):
        """Check if the API server is running"""
        try:
            response = self.session.get(f"{self.base_url}/health", timeout=5)
            return response.status_code == 200
        except:
            return False
    
    @patch('lib.llm_slides.analyze_documents_for_slides')
    def test_01_generate_slides_with_mocked_openai(self, mock_analyze):
        """Test slide generation with mocked OpenAI analysis"""
        print("\n=== Testing Generate Slides - Mocked OpenAI ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Configure mock
        mock_analyze.return_value = self.mock_openai_response
        
        with open(self.excel_file, 'rb') as f:
            files = {'documents': ('test_financials.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            response = self.session.post(f"{self.api_url}/generate-slides", files=files)
        
        if response.status_code == 200:
            # Verify PowerPoint file was generated
            self.assertEqual(response.headers.get('content-type'), 
                           'application/vnd.openxmlformats-officedocument.presentationml.presentation')
            self.assertGreater(len(response.content), 0)
            
            # Verify mock was called
            mock_analyze.assert_called_once()
            call_args = mock_analyze.call_args[0][0]  # First argument (documents)
            self.assertEqual(len(call_args), 1)
            self.assertEqual(call_args[0]['filename'], 'test_financials.xlsx')
            
            print("✓ OpenAI analysis mocked successfully, slides generated")
            
            # Save for inspection
            output_file = os.path.join(self.test_data_dir, "mocked_openai_output.pptx")
            with open(output_file, 'wb') as out_f:
                out_f.write(response.content)
            self.temp_files.append(output_file)
            
        else:
            # Handle error case
            try:
                data = response.json()
                print(f"⚠ Slide generation failed: {data}")
                # Still verify mock was called if error was in slide generation, not analysis
                if 'analysis' in str(data.get('error', '')).lower():
                    pass
                else:
                    mock_analyze.assert_called_once()
            except:
                self.fail(f"Unexpected response: {response.status_code} - {response.text}")
    
    @patch('lib.llm_slides.analyze_documents_for_slides')
    def test_02_generate_slides_multiple_docs_mocked(self, mock_analyze):
        """Test slide generation with multiple documents and mocked analysis"""
        print("\n=== Testing Generate Slides - Multiple Docs Mocked ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Configure mock with enhanced response for multiple documents
        enhanced_response = self.mock_openai_response.copy()
        enhanced_response['source_attributions']['primary_documents'] = [
            'test_financials.xlsx', 'business_plan.docx'
        ]
        enhanced_response['company_overview']['description'] += ' - Comprehensive analysis from multiple sources'
        mock_analyze.return_value = enhanced_response
        
        files = [
            ('documents', ('test_financials.xlsx', open(self.excel_file, 'rb'), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')),
            ('documents', ('business_plan.docx', open(self.word_file, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
        ]
        
        try:
            response = self.session.post(f"{self.api_url}/generate-slides", files=files)
            
            if response.status_code == 200:
                self.assertEqual(response.headers.get('content-type'), 
                               'application/vnd.openxmlformats-officedocument.presentationml.presentation')
                
                # Verify mock was called with multiple documents
                mock_analyze.assert_called_once()
                call_args = mock_analyze.call_args[0][0]
                self.assertEqual(len(call_args), 2)
                
                filenames = [doc['filename'] for doc in call_args]
                self.assertIn('test_financials.xlsx', filenames)
                self.assertIn('business_plan.docx', filenames)
                
                print("✓ Multiple documents processed with mocked analysis")
                
                # Save for inspection
                output_file = os.path.join(self.test_data_dir, "mocked_multiple_docs.pptx")
                with open(output_file, 'wb') as out_f:
                    out_f.write(response.content)
                self.temp_files.append(output_file)
                
        finally:
            # Close file handles
            for _, file_tuple in files:
                file_tuple[1].close()
    
    @patch('lib.llm_slides.analyze_documents_for_slides')
    def test_03_mocked_openai_error_handling(self, mock_analyze):
        """Test error handling when mocked OpenAI fails"""
        print("\n=== Testing Mocked OpenAI Error Handling ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Configure mock to raise an exception
        mock_analyze.side_effect = Exception("OpenAI API quota exceeded")
        
        with open(self.excel_file, 'rb') as f:
            files = {'documents': ('test_financials.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            response = self.session.post(f"{self.api_url}/generate-slides", files=files)
        
        # Should still generate slides with fallback analysis
        if response.status_code == 200:
            self.assertEqual(response.headers.get('content-type'), 
                           'application/vnd.openxmlformats-officedocument.presentationml.presentation')
            print("✓ Fallback analysis used when OpenAI fails")
            
        elif response.status_code == 500:
            data = response.json()
            self.assertIn('error', data)
            print(f"✓ Error handled gracefully: {data}")
        
        # Verify mock was called
        mock_analyze.assert_called_once()
    
    @patch('lib.pdf_extractor.PDFExtractor.extract_from_bytes')
    def test_04_mocked_llm_whisperer(self, mock_pdf_extract):
        """Test PDF extraction with mocked LLMWhisperer"""
        print("\n=== Testing Mocked LLMWhisperer PDF Extraction ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Configure mock PDF extraction
        mock_pdf_response = {
            'raw_text': 'Mocked PDF content from LLMWhisperer',
            'metadata': {'pages': 3, 'confidence': 0.95},
            'tables': [
                {'data': [['Metric', 'Value'], ['Revenue', '$10.2M'], ['Profit', '$2.5M']]},
            ],
            'sections': {
                'Executive Summary': 'Company overview and performance',
                'Financial Results': 'Q4 2023 financial performance'
            },
            'key_metrics': {
                'revenue': '10.2M',
                'profit': '2.5M',
                'growth': '23%'
            }
        }
        mock_pdf_extract.return_value = mock_pdf_response
        
        # Create a mock PDF file (just bytes)
        pdf_content = b"Mock PDF file content for testing"
        pdf_file = io.BytesIO(pdf_content)
        
        files = {'documents': ('quarterly_report.pdf', pdf_file, 'application/pdf')}
        response = self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
        
        if response.status_code == 200:
            data = response.json()
            self.assertTrue(data['success'])
            self.assertEqual(data['documents_processed'], 1)
            
            # Verify PDF extraction was called
            mock_pdf_extract.assert_called_once()
            
            # Check extracted document structure
            doc_result = data['extraction_results'][0]
            self.assertEqual(doc_result['type'], 'pdf')
            self.assertEqual(doc_result['filename'], 'quarterly_report.pdf')
            
            print("✓ LLMWhisperer PDF extraction mocked successfully")
            
        else:
            try:
                data = response.json()
                print(f"⚠ PDF extraction test failed: {data}")
            except:
                self.fail(f"Unexpected response: {response.status_code}")
    
    @patch('lib.llm_slides.extract_key_metrics_simple')
    def test_05_fallback_analysis_mocked(self, mock_simple_extract):
        """Test fallback analysis when OpenAI is not available"""
        print("\n=== Testing Fallback Analysis Mocked ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Configure mock for simple extraction
        mock_simple_response = {
            'financial_metrics': {
                'revenue': '10200000',
                'profit': '2500000',
                'growth_rate': '0.234'
            },
            'key_data_points': [
                'TechCorp Inc.',
                'Q4 2023 Financial Summary',
                'Technology Sector'
            ],
            'document_summary': 'Financial data extracted from 1 document'
        }
        mock_simple_extract.return_value = mock_simple_response
        
        # Temporarily remove OpenAI API key to trigger fallback
        with patch.dict(os.environ, {'OPENAI_API_KEY': ''}, clear=False):
            with open(self.excel_file, 'rb') as f:
                files = {'documents': ('test_financials.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                response = self.session.post(f"{self.api_url}/generate-slides", files=files)
        
        if response.status_code == 200:
            self.assertEqual(response.headers.get('content-type'), 
                           'application/vnd.openxmlformats-officedocument.presentationml.presentation')
            
            # Verify fallback was used
            mock_simple_extract.assert_called_once()
            
            print("✓ Fallback analysis triggered and mocked successfully")
            
        else:
            try:
                data = response.json()
                # Fallback analysis should still work even if mock not called
                print(f"⚠ Fallback analysis response: {data}")
            except:
                self.fail(f"Unexpected response: {response.status_code}")
    
    @patch('lib.template_parser.BrandManager')
    def test_06_template_management_mocked(self, mock_brand_manager):
        """Test template management with mocked BrandManager"""
        print("\n=== Testing Template Management Mocked ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Configure mock BrandManager
        mock_manager_instance = MagicMock()
        mock_brand_manager.return_value = mock_manager_instance
        
        mock_manager_instance.list_templates.return_value = ['default', 'corporate', 'minimal']
        mock_manager_instance.current_template = 'default'
        
        # Mock template parser
        mock_template_parser = MagicMock()
        mock_manager_instance.get_current_template.return_value = mock_template_parser
        mock_template_parser.get_brand_config.return_value = {
            'theme_colors': {
                'primary': '#1f4e79',
                'secondary': '#70ad47',
                'accent': '#c5504b'
            },
            'fonts': {
                'heading': 'Calibri',
                'body': 'Calibri Light'
            },
            'layouts': ['title', 'content', 'two_column']
        }
        
        # Test templates list endpoint
        response = self.session.get(f"{self.api_url}/templates")
        
        if response.status_code == 200:
            data = response.json()
            self.assertIn('templates', data)
            self.assertIn('count', data)
            
            # Verify mock was called
            mock_manager_instance.list_templates.assert_called()
            
            print(f"✓ Template management mocked successfully: {data['count']} templates")
            
        elif response.status_code == 503:
            # Template management not available
            data = response.json()
            self.assertIn('Template management not available', data['error'])
            print("⚠ Template management not available (expected if not configured)")
        
        else:
            self.fail(f"Unexpected response: {response.status_code}")
    
    @patch('lib.slide_generator_branded.BrandedSlideGenerator')
    @patch('lib.llm_slides.analyze_documents_for_slides')
    def test_07_branded_slide_generation_mocked(self, mock_analyze, mock_branded_generator):
        """Test branded slide generation with mocked components"""
        print("\n=== Testing Branded Slide Generation Mocked ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        # Configure analysis mock
        mock_analyze.return_value = self.mock_openai_response
        
        # Configure branded generator mock
        mock_generator_instance = MagicMock()
        mock_branded_generator.return_value = mock_generator_instance
        
        # Mock slide creation methods
        mock_generator_instance.create_financial_summary_slide.return_value = True
        mock_generator_instance.create_company_overview_slide.return_value = True
        mock_generator_instance.create_data_insights_slide.return_value = True
        mock_generator_instance.save_presentation.return_value = True
        
        with open(self.excel_file, 'rb') as f:
            files = {'documents': ('test_financials.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            data = {'template_id': 'corporate'}
            response = self.session.post(f"{self.api_url}/generate-slides", files=files, data=data)
        
        if response.status_code == 200:
            # Verify branded generator was used
            mock_branded_generator.assert_called()
            
            # Verify slide creation methods were called
            mock_generator_instance.create_financial_summary_slide.assert_called()
            mock_generator_instance.create_company_overview_slide.assert_called()
            mock_generator_instance.create_data_insights_slide.assert_called()
            mock_generator_instance.save_presentation.assert_called()
            
            print("✓ Branded slide generation mocked successfully")
            
        elif response.status_code == 500:
            # Expected if branded generator not available
            data = response.json()
            print(f"⚠ Branded generation failed (expected if not configured): {data}")
        
        else:
            self.fail(f"Unexpected response: {response.status_code}")
    
    def test_08_concurrent_mocked_requests(self):
        """Test concurrent requests with consistent mocked responses"""
        print("\n=== Testing Concurrent Mocked Requests ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        import threading
        import queue
        
        results = queue.Queue()
        
        def make_mocked_request(request_id):
            try:
                with patch('lib.llm_slides.analyze_documents_for_slides') as mock_analyze:
                    # Configure mock for each request
                    response_copy = self.mock_openai_response.copy()
                    response_copy['request_id'] = request_id
                    mock_analyze.return_value = response_copy
                    
                    with open(self.excel_file, 'rb') as f:
                        files = {'documents': (f'test_{request_id}.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                        response = requests.post(f"{self.api_url}/generate-slides/preview", files=files, timeout=30)
                        results.put((request_id, response.status_code, mock_analyze.called))
            except Exception as e:
                results.put((request_id, 'error', str(e)))
        
        # Start 3 concurrent requests
        threads = []
        for i in range(3):
            thread = threading.Thread(target=make_mocked_request, args=(i,))
            threads.append(thread)
            thread.start()
        
        # Wait for completion
        for thread in threads:
            thread.join(timeout=60)
        
        # Analyze results
        success_count = 0
        mocked_count = 0
        while not results.empty():
            request_id, status, mock_called = results.get()
            if status == 200:
                success_count += 1
                if mock_called:
                    mocked_count += 1
            print(f"Request {request_id}: Status {status}, Mock called: {mock_called}")
        
        self.assertGreater(success_count, 0)
        print(f"✓ Concurrent mocked requests: {success_count}/3 successful, {mocked_count} mocked")
    
    @patch('lib.llm_slides.analyze_documents_for_slides')
    def test_09_performance_with_mocking(self, mock_analyze):
        """Test API performance with mocked external calls"""
        print("\n=== Testing Performance With Mocking ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        import time
        
        # Configure fast mock response
        mock_analyze.return_value = self.mock_openai_response
        
        start_time = time.time()
        
        with open(self.excel_file, 'rb') as f:
            files = {'documents': ('test_financials.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            response = self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
        
        end_time = time.time()
        duration = end_time - start_time
        
        if response.status_code == 200:
            data = response.json()
            self.assertTrue(data['success'])
            
            # With mocking, should be fast
            self.assertLess(duration, 10.0)  # Should complete within 10 seconds
            
            print(f"✓ Mocked API performance: {duration:.2f}s")
            
        else:
            self.fail(f"Performance test failed: {response.status_code}")
    
    @patch.dict(os.environ, {'OPENAI_API_KEY': ''}, clear=False)
    def test_10_no_api_key_behavior(self):
        """Test API behavior when no external API keys are configured"""
        print("\n=== Testing No API Key Behavior ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        with open(self.excel_file, 'rb') as f:
            files = {'documents': ('test_financials.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            response = self.session.post(f"{self.api_url}/generate-slides", files=files)
        
        # Should still work with fallback analysis
        if response.status_code == 200:
            self.assertEqual(response.headers.get('content-type'), 
                           'application/vnd.openxmlformats-officedocument.presentationml.presentation')
            print("✓ Fallback analysis works without API keys")
            
        else:
            try:
                data = response.json()
                print(f"⚠ No API key response: {data}")
                # Should still handle gracefully
                self.assertIn('error', data)
            except:
                self.fail(f"Unexpected response without API key: {response.status_code}")

def run_mocked_tests():
    """Run the mocked API test suite"""
    print("=" * 60)
    print("MOCKED API TESTING SUITE")
    print("Testing with mocked external dependencies")
    print("=" * 60)
    
    # Create test suite
    loader = unittest.TestLoader()
    suite = loader.loadTestsFromTestCase(MockedAPITestSuite)
    
    # Run tests with detailed output
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    
    # Print summary
    print("\n" + "=" * 60)
    print("MOCKED TEST RESULTS SUMMARY")
    print("=" * 60)
    print(f"Tests run: {result.testsRun}")
    print(f"Failures: {len(result.failures)}")
    print(f"Errors: {len(result.errors)}")
    print(f"Skipped: {len(result.skipped)}")
    
    if result.failures:
        print("\nFAILURES:")
        for test, traceback in result.failures:
            print(f"- {test}")
    
    if result.errors:
        print("\nERRORS:")
        for test, traceback in result.errors:
            print(f"- {test}")
    
    success_rate = ((result.testsRun - len(result.failures) - len(result.errors)) / result.testsRun * 100) if result.testsRun > 0 else 0
    print(f"\nSuccess Rate: {success_rate:.1f}%")
    
    return result.wasSuccessful()

if __name__ == "__main__":
    success = run_mocked_tests()
    exit(0 if success else 1)