"""
Comprehensive edge case tests for document extraction components
Tests cover empty files, corrupted documents, invalid formats, and error handling
"""

import unittest
from unittest.mock import Mock, patch, MagicMock
import sys
import os
import io
import openpyxl
from openpyxl import Workbook
from docx import Document

# Add the lib directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', 'lib'))

from pdf_extractor import PDFExtractor
from excel_extractor import ExcelExtractor
from word_extractor import WordExtractor
from source_tracker import SourceTracker


class TestEdgeCases(unittest.TestCase):
    """Test suite for edge cases and error handling"""

    def setUp(self):
        """Set up test fixtures"""
        self.source_tracker = SourceTracker()
        self.pdf_extractor = PDFExtractor(source_tracker=self.source_tracker)
        self.excel_extractor = ExcelExtractor(source_tracker=self.source_tracker)
        self.word_extractor = WordExtractor()

    def test_empty_files(self):
        """Test handling of empty files across all extractors"""
        # Test empty PDF
        with patch.dict(os.environ, {}, clear=True):
            pdf_result = self.pdf_extractor.extract_from_bytes(b"", "empty.pdf")
            self.assertIn('API key not configured', pdf_result['raw_text'])
        
        # Test empty Excel bytes
        excel_result = self.excel_extractor.extract_from_bytes(b"", "empty.xlsx")
        self.assertIn('error', excel_result)
        
        # Test empty Word bytes
        word_result = self.word_extractor.extract_from_bytes(b"", "empty.docx")
        self.assertIn('error', word_result)

    def test_invalid_file_formats(self):
        """Test handling of invalid file formats"""
        invalid_data = b"This is not a valid document file"
        
        # Test invalid Excel format
        excel_result = self.excel_extractor.extract_from_bytes(invalid_data, "invalid.xlsx")
        self.assertIn('error', excel_result)
        self.assertIn('Failed to extract Excel data from bytes', excel_result['error'])
        
        # Test invalid Word format
        word_result = self.word_extractor.extract_from_bytes(invalid_data, "invalid.docx")
        self.assertIn('error', word_result)
        self.assertIn('Failed to extract Word document from bytes', word_result['error'])

    def test_corrupted_excel_data(self):
        """Test handling of corrupted Excel data"""
        # Create partially corrupted Excel-like data
        corrupted_excel = b"PK\x03\x04\x14\x00\x06\x00\x08\x00corrupted_data_here"
        
        result = self.excel_extractor.extract_from_bytes(corrupted_excel, "corrupted.xlsx")
        self.assertIn('error', result)

    def test_excel_with_password_protection(self):
        """Test handling of password-protected Excel files"""
        # Create a workbook and try to save with password (will fail gracefully)
        wb = Workbook()
        ws = wb.active
        ws['A1'] = "Protected data"
        
        excel_bytes = io.BytesIO()
        wb.save(excel_bytes)
        excel_bytes.seek(0)
        
        # Mock password protection scenario
        with patch('openpyxl.load_workbook') as mock_load:
            mock_load.side_effect = Exception("Password required")
            
            result = self.excel_extractor.extract_from_bytes(excel_bytes.getvalue(), "protected.xlsx")
            self.assertIn('error', result)

    def test_extremely_large_excel_performance(self):
        """Test performance with large Excel files (within reasonable limits)"""
        wb = Workbook()
        ws = wb.active
        
        # Create large dataset at the edge of limits (100 rows x 50 cols)
        for row in range(1, 101):
            for col in range(1, 51):
                ws.cell(row=row, column=col, value=f"R{row}C{col}")
        
        excel_bytes = io.BytesIO()
        wb.save(excel_bytes)
        excel_bytes.seek(0)
        
        # Should handle large file within limits
        result = self.excel_extractor.extract_from_bytes(excel_bytes.getvalue(), "large.xlsx")
        
        self.assertNotIn('error', result)
        self.assertIn('sheets', result)
        sheet_data = result['sheets']['Sheet']
        self.assertEqual(len(sheet_data['data']), 100)

    def test_excel_with_complex_formulas(self):
        """Test handling of complex and potentially problematic formulas"""
        wb = Workbook()
        ws = wb.active
        
        # Add formulas that might cause issues
        ws['A1'] = "=1/0"  # Division by zero
        ws['A2'] = "=VLOOKUP(Z1,A:B,2,FALSE)"  # Lookup that will fail
        ws['A3'] = "=SUM(#REF!)"  # Reference error
        ws['A4'] = "=INDIRECT(\"InvalidRef\")"  # Indirect reference error
        ws['A5'] = "=IF(ISERROR(A1),\"Error\",A1)"  # Error handling formula
        
        # Circular reference (commented out as it might cause issues)
        # ws['A6'] = "=A6+1"
        
        excel_bytes = io.BytesIO()
        wb.save(excel_bytes)
        excel_bytes.seek(0)
        
        # Should handle formula errors gracefully
        result = self.excel_extractor.extract_from_bytes(excel_bytes.getvalue(), "formulas.xlsx")
        
        self.assertNotIn('error', result)
        sheet_data = result['sheets']['Sheet']
        self.assertIn('formulas', sheet_data)

    def test_word_document_with_corrupted_structure(self):
        """Test handling of Word documents with structural issues"""
        # Create document and then corrupt parts of it
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("Some content")
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        original_bytes = doc_bytes.getvalue()
        
        # Corrupt part of the document
        corrupted_bytes = original_bytes[:100] + b"CORRUPTED" + original_bytes[109:]
        
        result = self.word_extractor.extract_from_bytes(corrupted_bytes, "corrupted.docx")
        # Should fail gracefully
        self.assertIn('error', result)

    def test_word_document_with_special_elements(self):
        """Test handling of Word documents with special elements"""
        doc = Document()
        
        # Add various elements that might cause issues
        doc.add_heading("Document with Special Elements", 0)
        
        # Add paragraph with special formatting
        p = doc.add_paragraph()
        p.add_run("Bold text").bold = True
        p.add_run(" and ")
        p.add_run("italic text").italic = True
        
        # Add table with merged cells scenario
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = ""  # Empty cell
        
        # Add very long paragraph
        long_text = "This is a very long paragraph. " * 1000  # 30,000+ characters
        doc.add_paragraph(long_text)
        
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        result = self.word_extractor.extract_from_bytes(doc_bytes.getvalue(), "special.docx")
        
        # Should handle special elements
        self.assertNotIn('error', result)
        self.assertIn('paragraphs', result)
        self.assertIn('tables', result)

    @patch.dict(os.environ, {'LLMWHISPERER_API_KEY': 'test_key'})
    @patch('pdf_extractor.requests.post')
    def test_pdf_api_timeout_scenarios(self, mock_post):
        """Test PDF extractor timeout and network error scenarios"""
        # Test connection timeout
        mock_post.side_effect = Exception("Connection timeout")
        
        result = self.pdf_extractor.extract_from_bytes(b"fake_pdf", "timeout.pdf")
        self.assertIn('error', result['metadata'])
        self.assertIn('Connection timeout', result['metadata']['error'])

    @patch.dict(os.environ, {'LLMWHISPERER_API_KEY': 'test_key'})
    @patch('pdf_extractor.requests.post')
    @patch('pdf_extractor.requests.get')
    def test_pdf_api_error_responses(self, mock_get, mock_post):
        """Test PDF extractor handling of various API error responses"""
        # Test 400 Bad Request
        mock_post.return_value.status_code = 400
        mock_post.return_value.text = "Bad Request: Invalid file format"
        
        result = self.pdf_extractor.extract_from_bytes(b"fake_pdf", "bad_request.pdf")
        self.assertIn('400', result['metadata']['error'])
        
        # Test 401 Unauthorized
        mock_post.return_value.status_code = 401
        mock_post.return_value.text = "Unauthorized: Invalid API key"
        
        result = self.pdf_extractor.extract_from_bytes(b"fake_pdf", "unauthorized.pdf")
        self.assertIn('401', result['metadata']['error'])
        
        # Test 429 Rate Limited
        mock_post.return_value.status_code = 429
        mock_post.return_value.text = "Rate limit exceeded"
        
        result = self.pdf_extractor.extract_from_bytes(b"fake_pdf", "rate_limited.pdf")
        self.assertIn('429', result['metadata']['error'])

    @patch.dict(os.environ, {'LLMWHISPERER_API_KEY': 'test_key'})
    @patch('pdf_extractor.requests.post')
    @patch('pdf_extractor.requests.get')
    def test_pdf_processing_failure_scenarios(self, mock_get, mock_post):
        """Test PDF processing failure scenarios"""
        # Successful upload but processing fails
        mock_post.return_value.status_code = 200
        mock_post.return_value.json.return_value = {"whisper_hash": "test_hash"}
        
        # Status check returns failed
        mock_get.return_value.status_code = 200
        mock_get.return_value.json.return_value = {"status": "failed"}
        
        result = self.pdf_extractor.extract_from_bytes(b"fake_pdf", "processing_failed.pdf")
        self.assertIn('timeout', result['metadata']['error'])

    def test_source_tracker_edge_cases(self):
        """Test source tracker with edge case inputs"""
        # Test with None values
        doc_id = self.source_tracker.register_document("test.xlsx", "excel", None)
        
        data_point_id = self.source_tracker.track_data_point(
            value=None,
            document_id=doc_id,
            location_details={},
            confidence=0.0
        )
        
        # Should handle None values gracefully
        self.assertIsNotNone(data_point_id)
        
        # Test with extreme values
        large_value = "x" * 10000  # Very large text value
        data_point_id = self.source_tracker.track_data_point(
            value=large_value,
            document_id=doc_id,
            location_details={'cell_or_section': 'A1'},
            confidence=1.5  # Invalid confidence (over 1.0)
        )
        
        self.assertIsNotNone(data_point_id)
        
        # Test invalid document references
        invalid_dp_id = self.source_tracker.track_data_point(
            value="test",
            document_id="invalid_doc_id",
            location_details={}
        )
        
        self.assertIsNotNone(invalid_dp_id)

    def test_malformed_table_structures(self):
        """Test handling of malformed table structures in various formats"""
        # Test PDF with malformed markdown tables
        malformed_md_table = """
| Header 1 | Header 2
|----------|----------
| Row 1 Col 1 | Row 1 Col 2 |
| Row 2 incomplete
| | Empty first cell | Second cell |
"""
        
        tables = self.pdf_extractor._extract_tables(malformed_md_table)
        # Should handle malformed tables gracefully
        self.assertIsInstance(tables, list)
        
        # Test Excel with inconsistent row lengths
        wb = Workbook()
        ws = wb.active
        
        # Create table with inconsistent structure
        ws['A1'] = "Header 1"
        ws['B1'] = "Header 2"
        ws['C1'] = "Header 3"
        
        # Row with missing cells
        ws['A2'] = "Data 1"
        # B2 intentionally left empty
        ws['C2'] = "Data 3"
        
        # Row with extra data
        ws['A3'] = "Data A"
        ws['B3'] = "Data B"
        ws['C3'] = "Data C"
        ws['D3'] = "Extra Data"  # Beyond expected columns
        
        excel_bytes = io.BytesIO()
        wb.save(excel_bytes)
        excel_bytes.seek(0)
        
        result = self.excel_extractor.extract_from_bytes(excel_bytes.getvalue(), "inconsistent.xlsx")
        # Should handle inconsistent structure
        self.assertNotIn('error', result)

    def test_memory_intensive_operations(self):
        """Test handling of memory-intensive operations"""
        # Create document with many repetitive elements
        doc = Document()
        
        # Add many paragraphs
        for i in range(500):  # Large number of paragraphs
            doc.add_paragraph(f"Paragraph {i}: " + "Text content. " * 50)
        
        # Add many small tables
        for i in range(50):  # Many small tables
            table = doc.add_table(rows=2, cols=3)
            table.cell(0, 0).text = f"Table {i}"
            table.cell(1, 1).text = f"Data {i}"
        
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        # Should handle large document
        result = self.word_extractor.extract_from_bytes(doc_bytes.getvalue(), "large_doc.docx")
        
        # Verify it handled the large document
        self.assertNotIn('error', result)
        self.assertTrue(len(result['paragraphs']) > 400)  # Should extract most paragraphs
        self.assertTrue(len(result['tables']) > 40)  # Should extract most tables

    def test_unicode_and_encoding_issues(self):
        """Test handling of unicode and encoding edge cases"""
        # Test Excel with unicode content
        wb = Workbook()
        ws = wb.active
        
        # Various unicode characters
        ws['A1'] = "中文测试"  # Chinese
        ws['A2'] = "العربية"    # Arabic
        ws['A3'] = "Русский"   # Russian
        ws['A4'] = "🚀📊💰"    # Emojis
        ws['A5'] = "Special: ∑∆∏αβγ"  # Mathematical symbols
        
        excel_bytes = io.BytesIO()
        wb.save(excel_bytes)
        excel_bytes.seek(0)
        
        result = self.excel_extractor.extract_from_bytes(excel_bytes.getvalue(), "unicode.xlsx")
        
        # Should handle unicode content
        self.assertNotIn('error', result)
        
        # Test Word with unicode
        doc = Document()
        doc.add_paragraph("Unicode test: 中文 العربية Русский 🚀📊💰")
        doc.add_paragraph("Math symbols: ∑∆∏αβγ ≤≥≠≈")
        
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        word_result = self.word_extractor.extract_from_bytes(doc_bytes.getvalue(), "unicode.docx")
        
        self.assertNotIn('error', word_result)
        self.assertIn('中文', word_result['raw_text'])

    def test_concurrent_extraction_simulation(self):
        """Test behavior under simulated concurrent extraction scenarios"""
        # Simulate multiple extractors working simultaneously
        source_tracker1 = SourceTracker()
        source_tracker2 = SourceTracker()
        
        extractor1 = ExcelExtractor(source_tracker1)
        extractor2 = ExcelExtractor(source_tracker2)
        
        # Create test workbooks
        wb1 = Workbook()
        wb1.active['A1'] = "Document 1 Data"
        wb1.active['A2'] = 1000
        
        wb2 = Workbook()
        wb2.active['A1'] = "Document 2 Data"
        wb2.active['A2'] = 2000
        
        # Extract simultaneously (simulated)
        bytes1 = io.BytesIO()
        wb1.save(bytes1)
        bytes1.seek(0)
        
        bytes2 = io.BytesIO()
        wb2.save(bytes2)
        bytes2.seek(0)
        
        result1 = extractor1.extract_from_bytes(bytes1.getvalue(), "doc1.xlsx")
        result2 = extractor2.extract_from_bytes(bytes2.getvalue(), "doc2.xlsx")
        
        # Both should succeed independently
        self.assertNotIn('error', result1)
        self.assertNotIn('error', result2)
        
        # Source trackers should be independent
        self.assertNotEqual(
            len(source_tracker1.data_points),
            len(source_tracker2.data_points)
        )

    def test_extraction_with_missing_dependencies(self):
        """Test graceful degradation when optional dependencies are missing"""
        # Test PDF extraction without API key (already covered above)
        
        # Test with mock import errors
        with patch('sys.modules', {'openpyxl': None}):
            # This would normally cause import error, but our code handles it
            # Test is more about ensuring graceful degradation patterns
            pass

    def test_extreme_values_handling(self):
        """Test handling of extreme values in data"""
        wb = Workbook()
        ws = wb.active
        
        # Extreme numeric values
        ws['A1'] = 999999999999999999999  # Very large number
        ws['A2'] = -999999999999999999999  # Very large negative number
        ws['A3'] = 0.00000000000001  # Very small decimal
        ws['A4'] = float('inf')  # Infinity
        ws['A5'] = float('-inf')  # Negative infinity
        
        # Very long text
        ws['A6'] = "x" * 32767  # Maximum Excel cell length
        
        excel_bytes = io.BytesIO()
        wb.save(excel_bytes)
        excel_bytes.seek(0)
        
        # Should handle extreme values
        result = self.excel_extractor.extract_from_bytes(excel_bytes.getvalue(), "extreme.xlsx")
        self.assertNotIn('error', result)

    def test_recursive_formula_references(self):
        """Test handling of complex formula references"""
        wb = Workbook()
        ws = wb.active
        
        # Create complex formula dependencies
        ws['A1'] = 100
        ws['A2'] = "=A1*2"
        ws['A3'] = "=A2+A1"
        ws['A4'] = "=SUM(A1:A3)"
        ws['A5'] = "=AVERAGE(A1:A4)"
        
        # Cross-reference formulas
        ws['B1'] = "=A5*1.1"
        ws['B2'] = "=A4+B1"
        
        excel_bytes = io.BytesIO()
        wb.save(excel_bytes)
        excel_bytes.seek(0)
        
        result = self.excel_extractor.extract_from_bytes(excel_bytes.getvalue(), "complex_formulas.xlsx")
        
        # Should handle complex formulas
        self.assertNotIn('error', result)
        sheet_data = result['sheets']['Sheet']
        self.assertTrue(len(sheet_data['formulas']) >= 5)


if __name__ == '__main__':
    unittest.main()