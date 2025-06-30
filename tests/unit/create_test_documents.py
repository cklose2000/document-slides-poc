"""
Create sample test documents for comprehensive testing
Generates Excel, Word documents with known data for validation
"""

import os
import sys
from io import BytesIO
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from docx import Document
from docx.shared import Inches

# Add the lib directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', 'lib'))


def create_sample_excel_financial():
    """Create comprehensive financial Excel document"""
    wb = Workbook()
    
    # Remove default sheet and create named sheets
    wb.remove(wb.active)
    
    # Financial Summary Sheet
    summary_sheet = wb.create_sheet("Financial Summary")
    
    # Headers
    summary_sheet['A1'] = "TechCorp Financial Summary 2023"
    summary_sheet['A1'].font = Font(bold=True, size=14)
    
    # Key Metrics
    summary_sheet['A3'] = "Key Financial Metrics"
    summary_sheet['A3'].font = Font(bold=True)
    
    metrics_data = [
        ['Metric', 'Value', 'Previous Year', 'Growth'],
        ['Total Revenue', 12500000, 9800000, '=B5-C5'],
        ['Gross Profit', 7500000, 5880000, '=B6-C6'],
        ['Net Income', 2250000, 1470000, '=B7-C7'],
        ['EBITDA', 3125000, 2156000, '=B8-C8'],
        ['Gross Margin', '=B6/B5', '=C6/C5', '=B9-C9'],
        ['Net Margin', '=B7/B5', '=C7/C5', '=B10-C10']
    ]
    
    for row_idx, row_data in enumerate(metrics_data, start=4):
        for col_idx, value in enumerate(row_data, start=1):
            cell = summary_sheet.cell(row=row_idx, column=col_idx, value=value)
            if row_idx == 4:  # Header row
                cell.font = Font(bold=True)
    
    # Quarterly Performance Sheet
    quarterly_sheet = wb.create_sheet("Quarterly Performance")
    
    quarterly_sheet['A1'] = "Quarterly Revenue Breakdown"
    quarterly_sheet['A1'].font = Font(bold=True, size=12)
    
    quarterly_data = [
        ['Quarter', 'Revenue', 'Costs', 'Profit', 'Margin'],
        ['Q1 2023', 3000000, 1800000, '=B3-C3', '=D3/B3'],
        ['Q2 2023', 3200000, 1920000, '=B4-C4', '=D4/B4'],
        ['Q3 2023', 3100000, 1860000, '=B5-C5', '=D5/B5'],
        ['Q4 2023', 3200000, 1920000, '=B6-C6', '=D6/B6'],
        ['Total', '=SUM(B3:B6)', '=SUM(C3:C6)', '=SUM(D3:D6)', '=D7/B7']
    ]
    
    for row_idx, row_data in enumerate(quarterly_data, start=2):
        for col_idx, value in enumerate(row_data, start=1):
            cell = quarterly_sheet.cell(row=row_idx, column=col_idx, value=value)
            if row_idx == 2:  # Header row
                cell.font = Font(bold=True)
    
    # Regional Performance Sheet
    regional_sheet = wb.create_sheet("Regional Performance")
    
    regional_sheet['A1'] = "Revenue by Region"
    regional_sheet['A1'].font = Font(bold=True, size=12)
    
    regional_data = [
        ['Region', '2023 Revenue', '2022 Revenue', 'Growth %', 'Market Share'],
        ['North America', 6250000, 4900000, '=(B3-C3)/C3', 0.15],
        ['Europe', 3750000, 2940000, '=(B4-C4)/C4', 0.12],
        ['Asia Pacific', 2500000, 1960000, '=(B5-C5)/C5', 0.08],
        ['Total', '=SUM(B3:B5)', '=SUM(C3:C5)', '=(B6-C6)/C6', '=AVERAGE(E3:E5)']
    ]
    
    for row_idx, row_data in enumerate(regional_data, start=2):
        for col_idx, value in enumerate(row_data, start=1):
            cell = regional_sheet.cell(row=row_idx, column=col_idx, value=value)
            if row_idx == 2:  # Header row
                cell.font = Font(bold=True)
    
    # KPIs Dashboard
    kpi_sheet = wb.create_sheet("KPIs Dashboard")
    
    kpi_sheet['A1'] = "Key Performance Indicators"
    kpi_sheet['A1'].font = Font(bold=True, size=14)
    
    # Reference other sheets
    kpi_data = [
        ['KPI', 'Value', 'Target', 'Status'],
        ['Revenue Growth', "='Financial Summary'!D5/C5", 0.25, '=IF(B3>=C3,"✓","✗")'],
        ['Gross Margin', "='Financial Summary'!B9", 0.60, '=IF(B4>=C4,"✓","✗")'],
        ['Net Margin', "='Financial Summary'!B10", 0.18, '=IF(B5>=C5,"✓","✗")'],
        ['Regional Diversity', "=STDEV('Regional Performance'!E3:E5)", 0.05, '=IF(B6<=C6,"✓","✗")']
    ]
    
    for row_idx, row_data in enumerate(kpi_data, start=3):
        for col_idx, value in enumerate(row_data, start=1):
            cell = kpi_sheet.cell(row=row_idx, column=col_idx, value=value)
            if row_idx == 3:  # Header row
                cell.font = Font(bold=True)
    
    return wb


def create_sample_word_business_plan():
    """Create comprehensive business plan Word document"""
    doc = Document()
    
    # Title Page
    title = doc.add_heading('TechCorp Business Plan 2024', 0)
    
    doc.add_paragraph('Confidential Business Plan')
    doc.add_paragraph('Prepared by: Executive Team')
    doc.add_paragraph('Date: January 2024')
    doc.add_paragraph('Version: 2.1')
    
    doc.add_page_break()
    
    # Executive Summary
    exec_summary = doc.add_heading('Executive Summary', level=1)
    
    doc.add_paragraph(
        'TechCorp is positioned for exceptional growth in 2024, building on our strong '
        'foundation of innovative products and expanding market presence. With projected '
        'revenue of $18.5M (48% growth), net income of $3.7M (64% growth), and strategic '
        'expansion into three new markets, we are well-positioned to capture significant '
        'market share in the rapidly evolving technology sector.'
    )
    
    doc.add_paragraph(
        'Key highlights include: successful product launch generating $2.1M in pre-orders, '
        'strategic partnership with GlobalTech Inc. worth $4.2M annually, and completion '
        'of Series B funding round raising $12M in growth capital.'
    )
    
    # Market Analysis
    market_heading = doc.add_heading('Market Analysis', level=1)
    
    doc.add_heading('Industry Overview', level=2)
    doc.add_paragraph(
        'The global technology solutions market is experiencing unprecedented growth, '
        'with our target segment expanding at 32% CAGR. Market size reached $125B in 2023 '
        'and is projected to reach $195B by 2026. Key drivers include digital transformation '
        'initiatives, cloud migration, and AI adoption across enterprises.'
    )
    
    doc.add_heading('Competitive Landscape', level=2)
    doc.add_paragraph(
        'Our primary competitors include TechGiant Corp (35% market share, $43B revenue), '
        'InnovateSoft Ltd (18% market share, $22B revenue), and emerging players like '
        'StartupTech Inc (3% market share, $3.8B revenue). Our competitive advantages '
        'include superior customer satisfaction (Net Promoter Score: 67), faster time-to-market '
        '(average 6 months vs industry 12 months), and lower total cost of ownership '
        '(40% savings vs competitors).'
    )
    
    # Financial Projections
    financial_heading = doc.add_heading('Financial Projections', level=1)
    
    doc.add_paragraph('Three-year financial outlook with detailed quarterly breakdowns:')
    
    # Financial table
    financial_table = doc.add_table(rows=1, cols=4)
    financial_table.style = 'Table Grid'
    
    # Headers
    hdr_cells = financial_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = '2024'
    hdr_cells[2].text = '2025'
    hdr_cells[3].text = '2026'
    
    # Financial data
    financial_data = [
        ('Revenue', '$18.5M', '$27.8M', '$41.7M'),
        ('Gross Revenue', '$11.1M', '$16.7M', '$25.0M'),
        ('Operating Expenses', '$8.9M', '$12.5M', '$17.2M'),
        ('Net Income', '$3.7M', '$6.8M', '$11.5M'),
        ('Cash Flow', '$5.2M', '$9.1M', '$14.8M'),
        ('Employees', '185', '275', '420'),
        ('Customers', '2,400', '4,200', '7,800')
    ]
    
    for metric, val_2024, val_2025, val_2026 in financial_data:
        row_cells = financial_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = val_2024
        row_cells[2].text = val_2025
        row_cells[3].text = val_2026
    
    # Product Strategy
    product_heading = doc.add_heading('Product Strategy', level=1)
    
    doc.add_heading('Current Product Portfolio', level=2)
    doc.add_paragraph(
        'Our flagship product TechSuite Pro generates $8.2M annually (66% of total revenue) '
        'with 94% customer satisfaction and 89% renewal rate. Supporting products include '
        'TechAnalytics ($2.1M revenue), TechConnect ($1.8M revenue), and TechSecure ($0.4M revenue).'
    )
    
    doc.add_heading('New Product Development', level=2)
    doc.add_paragraph(
        'Investment of $3.2M in R&D will deliver four new products: TechAI Platform (launch Q2 2024, '
        'projected $2.8M revenue), TechMobile Suite (launch Q3 2024, projected $1.5M revenue), '
        'TechIntegration Hub (launch Q4 2024, projected $1.1M revenue), and TechAdvanced Analytics '
        '(launch Q1 2025, projected $2.2M revenue).'
    )
    
    # Market Expansion
    expansion_heading = doc.add_heading('Market Expansion Strategy', level=1)
    
    doc.add_paragraph(
        'Geographic expansion into European Union (target: $4.5M revenue, 15 enterprise clients), '
        'Asia-Pacific region (target: $2.8M revenue, 25 SMB clients), and Latin America '
        '(target: $1.2M revenue, 35 small business clients). Total investment required: $2.1M '
        'across 18 months with break-even projected by month 14.'
    )
    
    # Risk Analysis
    risk_heading = doc.add_heading('Risk Analysis', level=1)
    
    doc.add_paragraph(
        'Primary risks include: Market competition intensification (High probability, High impact), '
        'Technology disruption by AI/ML advances (Medium probability, High impact), '
        'Economic downturn affecting enterprise spending (Medium probability, Medium impact), '
        'Key talent retention challenges (High probability, Medium impact), and '
        'Regulatory changes in data privacy (Low probability, High impact).'
    )
    
    doc.add_paragraph(
        'Mitigation strategies: Competitive differentiation through innovation ($1.8M R&D investment), '
        'Strategic partnerships and acquisitions ($5M war chest), Economic resilience through '
        'diversified customer base and flexible pricing, Talent retention program with equity '
        'incentives and professional development ($0.8M annual investment), and Proactive '
        'compliance program with legal counsel ($0.3M annual budget).'
    )
    
    # Implementation Timeline
    timeline_heading = doc.add_heading('Implementation Timeline', level=1)
    
    doc.add_paragraph('Phase 1 (Q1 2024): Product development completion, team expansion (25 new hires), market research finalization')
    doc.add_paragraph('Phase 2 (Q2 2024): TechAI Platform launch, European market entry, partnership agreements')
    doc.add_paragraph('Phase 3 (Q3 2024): TechMobile Suite launch, Asia-Pacific expansion, customer acquisition campaigns')
    doc.add_paragraph('Phase 4 (Q4 2024): TechIntegration Hub launch, Latin America entry, performance optimization')
    
    return doc


def create_sample_documents():
    """Create all sample test documents"""
    # Create sample_data directory if it doesn't exist
    sample_dir = os.path.join(os.path.dirname(__file__), '..', '..', 'sample_data')
    os.makedirs(sample_dir, exist_ok=True)
    
    print("📁 Creating comprehensive test documents...")
    
    # Create Excel financial document
    print("📊 Creating sample financial Excel document...")
    excel_wb = create_sample_excel_financial()
    excel_path = os.path.join(sample_dir, 'sample_financial_comprehensive.xlsx')
    excel_wb.save(excel_path)
    print(f"✅ Created: {excel_path}")
    
    # Create Word business plan
    print("📄 Creating sample business plan Word document...")
    word_doc = create_sample_word_business_plan()
    word_path = os.path.join(sample_dir, 'sample_business_plan_comprehensive.docx')
    word_doc.save(word_path)
    print(f"✅ Created: {word_path}")
    
    # Create simple Excel for basic testing
    print("📈 Creating simple Excel document...")
    simple_wb = Workbook()
    simple_ws = simple_wb.active
    simple_ws.title = "Simple Data"
    
    simple_data = [
        ['Item', 'Value', 'Category'],
        ['Revenue', 1000000, 'Financial'],
        ['Profit', 250000, 'Financial'],
        ['Growth Rate', '25%', 'Performance'],
        ['Customers', 150, 'Metrics'],
        ['Employee Count', 45, 'Metrics']
    ]
    
    for row_idx, row_data in enumerate(simple_data, start=1):
        for col_idx, value in enumerate(row_data, start=1):
            simple_ws.cell(row=row_idx, column=col_idx, value=value)
    
    simple_excel_path = os.path.join(sample_dir, 'sample_simple.xlsx')
    simple_wb.save(simple_excel_path)
    print(f"✅ Created: {simple_excel_path}")
    
    # Create simple Word document
    print("📝 Creating simple Word document...")
    simple_doc = Document()
    
    simple_doc.add_heading('Company Overview', 0)
    simple_doc.add_heading('Executive Summary', level=1)
    simple_doc.add_paragraph(
        'Our company achieved strong performance this year with revenue of $1,000,000 '
        'and profit of $250,000, representing 25% growth over the previous year.'
    )
    
    simple_doc.add_heading('Financial Performance', level=1)
    simple_doc.add_paragraph(
        'Key financial metrics show healthy growth: revenue increased to $1M, '
        'profit margins improved to 25%, and customer base expanded to 150 clients.'
    )
    
    # Simple table
    table = simple_doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    table_data = [
        ('Revenue', '$1,000,000'),
        ('Profit', '$250,000'),
        ('Growth', '25%'),
        ('Customers', '150')
    ]
    
    for metric, value in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    simple_word_path = os.path.join(sample_dir, 'sample_simple.docx')
    simple_doc.save(simple_word_path)
    print(f"✅ Created: {simple_word_path}")
    
    # Update README with document descriptions
    readme_path = os.path.join(sample_dir, 'README.txt')
    readme_content = """Sample Data for Document Extraction Testing
===========================================

This directory contains comprehensive test documents for validating 
document extraction functionality across PDF, Excel, and Word formats.

Files:
------

📊 sample_financial_comprehensive.xlsx
   - Multi-sheet Excel workbook with financial data
   - Contains formulas, cross-sheet references, and complex calculations
   - Sheets: Financial Summary, Quarterly Performance, Regional Performance, KPIs Dashboard
   - Tests: Formula parsing, cell referencing, source attribution

📄 sample_business_plan_comprehensive.docx  
   - Comprehensive business plan document
   - Contains headers, tables, financial projections, and structured content
   - Tests: Heading extraction, table parsing, paragraph structure, section identification

📈 sample_simple.xlsx
   - Basic Excel file with simple data structure
   - Tests: Basic extraction, cell values, simple formulas

📝 sample_simple.docx
   - Simple Word document with basic structure
   - Tests: Basic extraction, simple tables, heading detection

Test Data Values:
----------------

Excel Financial Data:
- Total Revenue: $12,500,000
- Gross Profit: $7,500,000  
- Net Income: $2,250,000
- EBITDA: $3,125,000
- Q1 Revenue: $3,000,000
- Q2 Revenue: $3,200,000
- Q3 Revenue: $3,100,000
- Q4 Revenue: $3,200,000

Word Business Plan Data:
- Projected 2024 Revenue: $18.5M
- Projected Net Income: $3.7M
- Market Size: $125B (2023)
- Growth Rate: 48%
- Employee Target: 185 (2024)

Usage:
------
These documents are used by the test suite to validate:
1. Accurate data extraction
2. Source attribution and tracking
3. Formula and calculation handling
4. Table and structure parsing
5. Cross-reference resolution
6. Error handling and edge cases

The test suite automatically loads these documents and validates that
extraction results match expected values and maintain proper source
attribution for slide generation.
"""
    
    with open(readme_path, 'w') as f:
        f.write(readme_content)
    
    print(f"✅ Updated: {readme_path}")
    print("\n🎉 All test documents created successfully!")
    print(f"📁 Documents saved to: {sample_dir}")


if __name__ == '__main__':
    create_sample_documents()