"""
Comprehensive tests for Word extractor with heading extraction, table parsing, and paragraph structure
Tests cover various document formats, styles, and edge cases
"""

import unittest
from unittest.mock import Mock, patch, MagicMock
import sys
import os
import io
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE

# Add the lib directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', 'lib'))

from word_extractor import WordExtractor


class TestWordExtractor(unittest.TestCase):
    """Test suite for Word extractor functionality"""

    def setUp(self):
        """Set up test fixtures"""
        self.word_extractor = WordExtractor()
        
    def create_test_document(self):
        """Create a test Word document for testing"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Business Analysis Report', 0)
        
        # Add executive summary
        exec_summary = doc.add_heading('Executive Summary', level=1)
        exec_para = doc.add_paragraph(
            'Our company has shown exceptional growth this year with revenue '
            'increasing by 25% to $1.2M. The market expansion strategy has '
            'proven successful across all regions.'
        )
        
        # Add financial performance section
        financial_heading = doc.add_heading('Financial Performance', level=1)
        financial_para = doc.add_paragraph(
            'Revenue for 2023 reached $1,200,000, representing a significant '
            'increase from the previous year. Profit margins improved to 18%, '
            'and EBITDA grew by 30% to $350,000.'
        )
        
        # Add a table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = '2023'
        hdr_cells[2].text = '2022'
        
        # Add table data
        metrics_data = [
            ('Revenue', '$1,200,000', '$960,000'),
            ('Profit', '$216,000', '$150,000'),
            ('EBITDA', '$350,000', '$270,000')
        ]
        
        for metric, value_2023, value_2022 in metrics_data:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value_2023
            row_cells[2].text = value_2022
        
        # Add growth section
        growth_heading = doc.add_heading('Growth Trends', level=1)
        growth_para = doc.add_paragraph(
            'The company experienced strong growth across all key metrics. '
            'Customer acquisition increased by 40%, and retention rates '
            'improved to 92%.'
        )
        
        # Add subsection
        growth_subsection = doc.add_heading('Market Expansion', level=2)
        market_para = doc.add_paragraph(
            'Our expansion into the European market contributed $300,000 in '
            'new revenue, representing 25% of total growth.'
        )
        
        # Add operations section
        ops_heading = doc.add_heading('Operations', level=1)
        ops_para = doc.add_paragraph(
            'Operational efficiency improvements reduced costs by 12% while '
            'maintaining service quality standards.'
        )
        
        return doc

    def test_extract_from_bytes_success(self):
        """Test successful Word document extraction from bytes"""
        # Create test document
        doc = self.create_test_document()
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        result = self.word_extractor.extract_from_bytes(doc_bytes.getvalue(), "test.docx")
        
        # Verify result structure
        self.assertIn('filename', result)
        self.assertIn('paragraphs', result)
        self.assertIn('tables', result)
        self.assertIn('headers', result)
        self.assertIn('key_sections', result)
        self.assertIn('raw_text', result)
        self.assertEqual(result['filename'], "test.docx")
        
        # Verify headers were extracted
        self.assertTrue(len(result['headers']) > 0)
        header_texts = [h['text'] for h in result['headers']]
        self.assertIn('Executive Summary', header_texts)
        self.assertIn('Financial Performance', header_texts)
        self.assertIn('Growth Trends', header_texts)
        
        # Verify tables were extracted
        self.assertTrue(len(result['tables']) > 0)
        table = result['tables'][0]
        self.assertIn('headers', table)
        self.assertIn('Metric', table['headers'])
        self.assertIn('2023', table['headers'])
        
        # Verify key sections were identified
        self.assertIn('financial_performance', result['key_sections'])
        self.assertIn('executive_summary', result['key_sections'])

    def test_extract_from_bytes_error_handling(self):
        """Test error handling for invalid Word data"""
        invalid_bytes = b"not_a_word_document"
        result = self.word_extractor.extract_from_bytes(invalid_bytes, "invalid.docx")
        
        self.assertIn('error', result)
        self.assertIn('Failed to extract Word document from bytes', result['error'])

    def test_extract_with_structure_file_path(self):
        """Test extraction using file path method"""
        # Create temporary document
        doc = self.create_test_document()
        temp_path = "/tmp/test_doc.docx"
        
        # Mock file operations
        with patch('word_extractor.Document') as mock_document:
            mock_document.return_value = doc
            
            result = self.word_extractor.extract_with_structure(temp_path)
            
            # Verify extraction was attempted
            mock_document.assert_called_once_with(temp_path)
            
            # Verify result structure
            self.assertIn('paragraphs', result)
            self.assertIn('tables', result)
            self.assertIn('headers', result)

    def test_is_header_style(self):
        """Test header style detection"""
        doc = Document()
        
        # Create paragraphs with different styles
        normal_para = doc.add_paragraph("Normal text")
        heading_para = doc.add_heading("This is a heading", level=1)
        title_para = doc.add_paragraph("Title text")
        title_para.style = 'Title'
        
        # Test header detection
        self.assertFalse(self.word_extractor._is_header_style(normal_para))
        self.assertTrue(self.word_extractor._is_header_style(heading_para))
        self.assertTrue(self.word_extractor._is_header_style(title_para))

    def test_get_header_level(self):
        """Test header level extraction"""
        doc = Document()
        
        # Create headers of different levels
        h1 = doc.add_heading("Level 1 Heading", level=1)
        h2 = doc.add_heading("Level 2 Heading", level=2)
        h3 = doc.add_heading("Level 3 Heading", level=3)
        
        self.assertEqual(self.word_extractor._get_header_level(h1), 1)
        self.assertEqual(self.word_extractor._get_header_level(h2), 2)
        self.assertEqual(self.word_extractor._get_header_level(h3), 3)

    def test_extract_table_comprehensive(self):
        """Test comprehensive table extraction"""
        doc = Document()
        
        # Create table with headers and data
        table = doc.add_table(rows=1, cols=4)
        
        # Headers
        headers = ['Product', 'Q1 Sales', 'Q2 Sales', 'Total']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # Data rows
        data_rows = [
            ('Product A', '$100,000', '$120,000', '$220,000'),
            ('Product B', '$85,000', '$95,000', '$180,000'),
            ('Product C', '$200,000', '$250,000', '$450,000')
        ]
        
        for row_data in data_rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = cell_data
        
        table_data = self.word_extractor._extract_table(table, 1)
        
        # Verify table structure
        self.assertEqual(table_data['id'], 'table_1')
        self.assertEqual(len(table_data['headers']), 4)
        self.assertEqual(len(table_data['rows']), 4)  # Headers + 3 data rows
        self.assertEqual(table_data['cell_count'], 16)  # 4 cols × 4 rows
        
        # Verify headers
        self.assertEqual(table_data['headers'], headers)
        
        # Verify cell positions
        first_data_row = table_data['rows'][1]  # Skip header row
        self.assertEqual(first_data_row[0]['position'], 'A2')
        self.assertEqual(first_data_row[1]['position'], 'B2')
        
        # Verify header detection
        header_row = table_data['rows'][0]
        for cell in header_row:
            self.assertTrue(cell['is_header'])
        
        data_row = table_data['rows'][1]
        for cell in data_row:
            self.assertFalse(cell['is_header'])

    def test_identify_key_sections(self):
        """Test identification of key document sections"""
        # Create mock paragraphs
        paragraphs = [
            {'id': 'para_1', 'text': 'Executive Summary', 'is_header': True},
            {'id': 'para_2', 'text': 'This document provides an overview of company performance.', 'is_header': False},
            {'id': 'para_3', 'text': 'Financial Performance Analysis', 'is_header': True},
            {'id': 'para_4', 'text': 'Revenue increased significantly this year.', 'is_header': False},
            {'id': 'para_5', 'text': 'Growth Strategy', 'is_header': True},
            {'id': 'para_6', 'text': 'Our expansion plans focus on market penetration.', 'is_header': False},
            {'id': 'para_7', 'text': 'Market Analysis', 'is_header': True},
            {'id': 'para_8', 'text': 'Industry trends show positive indicators.', 'is_header': False},
            {'id': 'para_9', 'text': 'Operations Overview', 'is_header': True},
            {'id': 'para_10', 'text': 'Business processes have been streamlined.', 'is_header': False}
        ]
        
        sections = self.word_extractor._identify_key_sections(paragraphs)
        
        # Verify key sections were identified
        self.assertIn('executive_summary', sections)
        self.assertIn('financial_performance', sections)
        self.assertIn('growth', sections)
        self.assertIn('market', sections)
        self.assertIn('operations', sections)
        
        # Verify section content
        exec_section = sections['executive_summary']
        self.assertEqual(len(exec_section), 2)  # Header + content
        self.assertEqual(exec_section[0]['type'], 'header')
        self.assertEqual(exec_section[1]['type'], 'content')

    def test_paragraph_processing(self):
        """Test paragraph processing and structure preservation"""
        doc = self.create_test_document()
        result = self.word_extractor._process_document(doc)
        
        # Verify paragraphs were processed
        self.assertTrue(len(result['paragraphs']) > 0)
        
        # Check paragraph structure
        for para in result['paragraphs']:
            self.assertIn('id', para)
            self.assertIn('text', para)
            self.assertIn('style', para)
            self.assertIn('is_header', para)
            self.assertTrue(para['id'].startswith('para_'))
        
        # Verify headers were identified in paragraphs
        header_paras = [p for p in result['paragraphs'] if p['is_header']]
        self.assertTrue(len(header_paras) > 0)
        
        # Verify raw text compilation
        self.assertIn('Business Analysis Report', result['raw_text'])
        self.assertIn('Executive Summary', result['raw_text'])
        self.assertIn('$1,200,000', result['raw_text'])

    def test_multiple_table_extraction(self):
        """Test extraction of multiple tables"""
        doc = Document()
        
        # Add first table - Financial Summary
        table1 = doc.add_table(rows=1, cols=3)
        hdr1 = table1.rows[0].cells
        hdr1[0].text = 'Metric'
        hdr1[1].text = 'Value'
        hdr1[2].text = 'Growth'
        
        row1 = table1.add_row().cells
        row1[0].text = 'Revenue'
        row1[1].text = '$1M'
        row1[2].text = '25%'
        
        # Add some text between tables
        doc.add_paragraph("Analysis of quarterly performance:")
        
        # Add second table - Quarterly Data
        table2 = doc.add_table(rows=1, cols=5)
        hdr2 = table2.rows[0].cells
        quarters = ['Quarter', 'Q1', 'Q2', 'Q3', 'Q4']
        for i, quarter in enumerate(quarters):
            hdr2[i].text = quarter
        
        row2 = table2.add_row().cells
        row2[0].text = 'Sales'
        quarterly_sales = ['250K', '300K', '280K', '320K']
        for i, sales in enumerate(quarterly_sales):
            row2[i+1].text = sales
        
        result = self.word_extractor._process_document(doc)
        
        # Verify both tables were extracted
        self.assertEqual(len(result['tables']), 2)
        
        # Verify first table
        table1_data = result['tables'][0]
        self.assertEqual(table1_data['id'], 'table_1')
        self.assertIn('Revenue', table1_data['headers'][0])
        
        # Verify second table
        table2_data = result['tables'][1]
        self.assertEqual(table2_data['id'], 'table_2')
        self.assertIn('Quarter', table2_data['headers'][0])

    def test_complex_document_structure(self):
        """Test extraction from complex document with nested headings"""
        doc = Document()
        
        # Create complex structure
        doc.add_heading('Annual Report 2023', 0)  # Title
        doc.add_paragraph('This is the annual report introduction.')
        
        # Level 1 sections
        doc.add_heading('Financial Overview', 1)
        doc.add_paragraph('Overall financial performance was strong.')
        
        # Level 2 subsections
        doc.add_heading('Revenue Analysis', 2)
        doc.add_paragraph('Revenue grew by 23% to $10.2M.')
        
        doc.add_heading('Cost Management', 2)
        doc.add_paragraph('Costs were controlled effectively.')
        
        # Level 3 subsections
        doc.add_heading('Direct Costs', 3)
        doc.add_paragraph('Direct costs decreased by 5%.')
        
        doc.add_heading('Indirect Costs', 3)
        doc.add_paragraph('Indirect costs remained stable.')
        
        # Another Level 1 section
        doc.add_heading('Market Position', 1)
        doc.add_paragraph('Market share increased to 15%.')
        
        result = self.word_extractor._process_document(doc)
        
        # Verify header levels were captured
        headers = result['headers']
        self.assertTrue(len(headers) >= 6)
        
        # Check specific header levels
        header_levels = {h['text']: h['level'] for h in headers}
        self.assertEqual(header_levels.get('Financial Overview'), 1)
        self.assertEqual(header_levels.get('Revenue Analysis'), 2)
        self.assertEqual(header_levels.get('Direct Costs'), 3)

    def test_edge_case_empty_document(self):
        """Test handling of empty document"""
        doc = Document()
        result = self.word_extractor._process_document(doc)
        
        # Should handle empty document gracefully
        self.assertEqual(len(result['paragraphs']), 0)
        self.assertEqual(len(result['tables']), 0)
        self.assertEqual(len(result['headers']), 0)
        self.assertEqual(result['raw_text'], '')

    def test_edge_case_document_with_only_tables(self):
        """Test handling of document containing only tables"""
        doc = Document()
        
        # Add only a table, no text
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Header 1'
        table.cell(0, 1).text = 'Header 2'
        table.cell(1, 0).text = 'Data 1'
        table.cell(1, 1).text = 'Data 2'
        
        result = self.word_extractor._process_document(doc)
        
        # Should extract table even without text
        self.assertEqual(len(result['tables']), 1)
        self.assertEqual(len(result['paragraphs']), 0)
        self.assertEqual(result['raw_text'], '')

    def test_edge_case_document_with_only_headers(self):
        """Test handling of document with only headers, no content"""
        doc = Document()
        
        doc.add_heading('Header 1', 1)
        doc.add_heading('Header 2', 1)
        doc.add_heading('Header 3', 2)
        
        result = self.word_extractor._process_document(doc)
        
        # Should identify headers
        self.assertEqual(len(result['headers']), 3)
        self.assertEqual(len(result['paragraphs']), 3)  # Headers are also paragraphs
        
        # All paragraphs should be headers
        for para in result['paragraphs']:
            self.assertTrue(para['is_header'])

    def test_text_with_special_characters(self):
        """Test handling of text with special characters and formatting"""
        doc = Document()
        
        # Add text with special characters
        doc.add_paragraph('Revenue: $1,234,567.89 (€1,000,000)')
        doc.add_paragraph('Growth: 25% → 30% (↑5%)')
        doc.add_paragraph('Symbols: α, β, γ, ∑, ∆, ∏')
        doc.add_paragraph('Unicode: 中文, العربية, Русский')
        
        result = self.word_extractor._process_document(doc)
        
        # Should handle special characters
        raw_text = result['raw_text']
        self.assertIn('$1,234,567.89', raw_text)
        self.assertIn('€1,000,000', raw_text)
        self.assertIn('25%', raw_text)
        self.assertIn('→', raw_text)
        self.assertIn('∑', raw_text)

    def test_table_with_merged_cells(self):
        """Test handling of tables with complex structure"""
        doc = Document()
        
        # Create table with various cell contents
        table = doc.add_table(rows=3, cols=3)
        
        # Headers
        table.cell(0, 0).text = 'Category'
        table.cell(0, 1).text = 'Value'
        table.cell(0, 2).text = 'Notes'
        
        # Data with empty cells
        table.cell(1, 0).text = 'Revenue'
        table.cell(1, 1).text = '$1,000,000'
        table.cell(1, 2).text = ''  # Empty cell
        
        table.cell(2, 0).text = ''  # Empty cell
        table.cell(2, 1).text = '$500,000'
        table.cell(2, 2).text = 'Projected'
        
        table_data = self.word_extractor._extract_table(table, 1)
        
        # Should handle empty cells gracefully
        self.assertEqual(len(table_data['rows']), 3)
        self.assertEqual(table_data['cell_count'], 9)
        
        # Verify empty cells are handled
        row2 = table_data['rows'][1]
        self.assertEqual(row2[2]['text'], '')  # Empty notes cell
        
        row3 = table_data['rows'][2]
        self.assertEqual(row3[0]['text'], '')  # Empty category cell

    def test_section_keyword_matching(self):
        """Test section identification with various keyword patterns"""
        paragraphs = [
            {'id': 'para_1', 'text': 'EXECUTIVE SUMMARY', 'is_header': True},
            {'id': 'para_2', 'text': 'Company overview content.', 'is_header': False},
            {'id': 'para_3', 'text': 'Revenue Performance', 'is_header': True},
            {'id': 'para_4', 'text': 'Financial metrics improved.', 'is_header': False},
            {'id': 'para_5', 'text': 'Business Growth Analysis', 'is_header': True},
            {'id': 'para_6', 'text': 'Expansion strategies discussed.', 'is_header': False},
            {'id': 'para_7', 'text': 'Industry Market Trends', 'is_header': True},
            {'id': 'para_8', 'text': 'Competitive positioning.', 'is_header': False},
            {'id': 'para_9', 'text': 'Operational Strategy', 'is_header': True},
            {'id': 'para_10', 'text': 'Process improvements.', 'is_header': False}
        ]
        
        sections = self.word_extractor._identify_key_sections(paragraphs)
        
        # Verify different keyword patterns are matched
        self.assertIn('executive_summary', sections)  # "executive summary"
        self.assertIn('financial_performance', sections)  # "revenue" -> financial
        self.assertIn('growth', sections)  # "growth"
        self.assertIn('market', sections)  # "market"
        self.assertIn('operations', sections)  # "operational" -> operations


class TestWordExtractorIntegration(unittest.TestCase):
    """Integration tests for Word extractor with real-world scenarios"""

    def setUp(self):
        """Set up integration test fixtures"""
        self.word_extractor = WordExtractor()

    def test_business_plan_extraction(self):
        """Test extraction from comprehensive business plan document"""
        doc = Document()
        
        # Title page
        doc.add_heading('TechCorp Business Plan 2024', 0)
        doc.add_paragraph('Prepared by: Management Team')
        doc.add_paragraph('Date: January 2024')
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(
            'TechCorp is positioned for significant growth in 2024 with projected '
            'revenue of $15M, representing 50% growth over 2023. Our innovative '
            'products and expanding market presence provide strong fundamentals '
            'for continued success.'
        )
        
        # Market Analysis
        doc.add_heading('Market Analysis', 1)
        doc.add_paragraph(
            'The technology sector shows robust growth with our target market '
            'expanding at 25% annually. Competitive analysis indicates strong '
            'positioning for market share capture.'
        )
        
        # Financial Projections Table
        doc.add_heading('Financial Projections', 1)
        doc.add_paragraph('Three-year financial outlook:')
        
        fin_table = doc.add_table(rows=1, cols=4)
        fin_headers = ['Metric', '2024', '2025', '2026']
        for i, header in enumerate(fin_headers):
            fin_table.cell(0, i).text = header
        
        financial_data = [
            ('Revenue', '$15.0M', '$22.5M', '$33.8M'),
            ('Gross Profit', '$9.0M', '$13.5M', '$20.3M'),
            ('Net Income', '$3.0M', '$5.6M', '$9.5M'),
            ('Employees', '150', '225', '340')
        ]
        
        for row_data in financial_data:
            row_cells = fin_table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = cell_data
        
        # Product Strategy
        doc.add_heading('Product Strategy', 1)
        doc.add_heading('Current Products', 2)
        doc.add_paragraph(
            'Our flagship product line generates $10M annually with 85% '
            'customer satisfaction. New features planned for Q2 release.'
        )
        
        doc.add_heading('New Product Development', 2)
        doc.add_paragraph(
            'Investment of $2M in R&D will yield 3 new products by end of 2024, '
            'targeting $5M additional revenue.'
        )
        
        # Risk Analysis
        doc.add_heading('Risk Analysis', 1)
        doc.add_paragraph(
            'Key risks include market competition (High), regulatory changes '
            '(Medium), and technology disruption (Medium). Mitigation strategies '
            'are in place for all identified risks.'
        )
        
        result = self.word_extractor._process_document(doc)
        
        # Verify comprehensive extraction
        self.assertTrue(len(result['paragraphs']) > 10)
        self.assertTrue(len(result['headers']) >= 7)
        self.assertEqual(len(result['tables']), 1)
        
        # Verify key sections were identified
        sections = result['key_sections']
        self.assertIn('executive_summary', sections)
        self.assertIn('market', sections)
        self.assertIn('financial_performance', sections)
        
        # Verify financial table extraction
        table = result['tables'][0]
        self.assertIn('Revenue', [row[0]['text'] for row in table['rows'][1:]])
        self.assertIn('$15.0M', table['rows'][1][1]['text'])
        
        # Verify nested heading structure
        header_levels = {h['text']: h['level'] for h in result['headers']}
        self.assertEqual(header_levels['Product Strategy'], 1)
        self.assertEqual(header_levels['Current Products'], 2)

    def test_annual_report_extraction(self):
        """Test extraction from annual report with complex structure"""
        doc = Document()
        
        # Cover page
        doc.add_heading('Annual Report 2023', 0)
        doc.add_heading('Building Tomorrow Today', level=1)
        
        # CEO Letter
        doc.add_heading('Letter from the CEO', 1)
        doc.add_paragraph(
            'Dear Shareholders, 2023 was a transformational year for our company. '
            'We achieved record revenue of $50M, up 35% from 2022, while expanding '
            'our global footprint to 12 countries.'
        )
        
        # Financial Highlights
        doc.add_heading('Financial Highlights', 1)
        
        # Key metrics table
        metrics_table = doc.add_table(rows=1, cols=3)
        metrics_headers = ['Metric', '2023', '2022']
        for i, header in enumerate(metrics_headers):
            metrics_table.cell(0, i).text = header
        
        key_metrics = [
            ('Total Revenue', '$50.2M', '$37.1M'),
            ('Net Income', '$8.4M', '$5.2M'),
            ('Earnings per Share', '$2.45', '$1.58'),
            ('Return on Equity', '18.5%', '14.2%'),
            ('Total Assets', '$125M', '$98M')
        ]
        
        for metric_data in key_metrics:
            row = metrics_table.add_row().cells
            for i, data in enumerate(metric_data):
                row[i].text = data
        
        # Business Segments
        doc.add_heading('Business Segments', 1)
        
        doc.add_heading('Technology Solutions', 2)
        doc.add_paragraph(
            'Our technology division generated $30M in revenue, representing '
            '60% of total company revenue. Customer retention rate of 95% '
            'demonstrates strong product-market fit.'
        )
        
        doc.add_heading('Professional Services', 2)
        doc.add_paragraph(
            'Services division contributed $20.2M with 40% gross margins. '
            'Expansion into consulting services drove 25% growth.'
        )
        
        # Geographic Performance
        doc.add_heading('Geographic Performance', 1)
        
        # Regional table
        regional_table = doc.add_table(rows=1, cols=4)
        regional_headers = ['Region', 'Revenue', 'Growth', 'Market Share']
        for i, header in enumerate(regional_headers):
            regional_table.cell(0, i).text = header
        
        regional_data = [
            ('North America', '$30M', '25%', '12%'),
            ('Europe', '$15M', '45%', '8%'),
            ('Asia Pacific', '$5.2M', '80%', '3%')
        ]
        
        for region_data in regional_data:
            row = regional_table.add_row().cells
            for i, data in enumerate(region_data):
                row[i].text = data
        
        # Future Outlook
        doc.add_heading('2024 Outlook', 1)
        doc.add_paragraph(
            'We project continued strong growth with revenue targets of $65-70M '
            'in 2024. Key initiatives include product line expansion, geographic '
            'diversification, and strategic acquisitions.'
        )
        
        result = self.word_extractor._process_document(doc)
        
        # Verify comprehensive document processing
        self.assertTrue(len(result['paragraphs']) > 15)
        self.assertTrue(len(result['headers']) >= 8)
        self.assertEqual(len(result['tables']), 2)
        
        # Verify financial data extraction
        tables = result['tables']
        
        # First table should be financial metrics
        fin_table = tables[0]
        self.assertIn('Total Revenue', [row[0]['text'] for row in fin_table['rows'][1:]])
        
        # Second table should be regional data
        regional_table = tables[1]
        self.assertIn('Region', regional_table['headers'])
        
        # Verify section identification
        sections = result['key_sections']
        self.assertTrue(len(sections) >= 2)
        
        # Verify raw text contains key financial figures
        raw_text = result['raw_text']
        self.assertIn('$50.2M', raw_text)
        self.assertIn('35%', raw_text)
        self.assertIn('$8.4M', raw_text)

    def test_technical_document_with_formulas(self):
        """Test extraction from technical document with equations and data"""
        doc = Document()
        
        # Title
        doc.add_heading('Financial Analysis Methodology', 0)
        
        # Introduction
        doc.add_heading('Overview', 1)
        doc.add_paragraph(
            'This document outlines the financial analysis methodology used '
            'for quarterly reporting and strategic planning.'
        )
        
        # Metrics section
        doc.add_heading('Key Performance Indicators', 1)
        doc.add_paragraph(
            'Primary KPIs include: Revenue Growth Rate = (Current Period - Previous Period) / Previous Period × 100'
        )
        doc.add_paragraph(
            'Return on Investment (ROI) = (Gain from Investment - Cost of Investment) / Cost of Investment × 100'
        )
        doc.add_paragraph(
            'Customer Acquisition Cost (CAC) = Total Sales & Marketing Expenses / Number of New Customers'
        )
        
        # Sample calculations table
        doc.add_heading('Sample Calculations', 1)
        calc_table = doc.add_table(rows=1, cols=4)
        calc_headers = ['Metric', 'Formula', 'Values', 'Result']
        for i, header in enumerate(calc_headers):
            calc_table.cell(0, i).text = header
        
        calculations = [
            ('Revenue Growth', '(1.2M - 1.0M) / 1.0M × 100', '1.2M, 1.0M', '20%'),
            ('Gross Margin', 'Gross Profit / Revenue × 100', '600K / 1.2M', '50%'),
            ('EBITDA Margin', 'EBITDA / Revenue × 100', '360K / 1.2M', '30%')
        ]
        
        for calc_data in calculations:
            row = calc_table.add_row().cells
            for i, data in enumerate(calc_data):
                row[i].text = data
        
        result = self.word_extractor._process_document(doc)
        
        # Verify extraction of technical content
        raw_text = result['raw_text']
        self.assertIn('Revenue Growth Rate', raw_text)
        self.assertIn('ROI', raw_text)
        self.assertIn('× 100', raw_text)
        
        # Verify table with formulas
        table = result['tables'][0]
        self.assertIn('Formula', table['headers'])
        self.assertTrue(any('1.2M - 1.0M' in row[2]['text'] for row in table['rows'][1:]))


if __name__ == '__main__':
    unittest.main()