from docx import Document
import io
import re
from typing import Dict, Any, Optional

try:
    from .source_tracker import SourceTracker
except ImportError:
    from source_tracker import SourceTracker

class WordExtractor:
    def __init__(self, source_tracker: Optional[SourceTracker] = None):
        """Initialize with optional source tracker for enhanced attribution"""
        self.source_tracker = source_tracker
    
    def extract_with_structure(self, file_path):
        """Extract maintaining paragraph numbers and structure"""
        try:
            doc = Document(file_path)
            return self._process_document(doc, file_path)
        except Exception as e:
            return {'error': f'Failed to extract Word document: {str(e)}'}
    
    def extract_from_bytes(self, word_bytes, filename="document.docx"):
        """Extract from bytes similar to other extractors"""
        try:
            word_file = io.BytesIO(word_bytes)
            doc = Document(word_file)
            result = self._process_document(doc, filename)
            result['filename'] = filename
            return result
        except Exception as e:
            return {'error': f'Failed to extract Word document from bytes: {str(e)}'}
    
    def _process_document(self, doc, file_path: str = "document.docx"):
        """Process the document and extract structured content with source tracking"""
        # Register document with source tracker if available
        document_id = None
        if self.source_tracker:
            document_id = self.source_tracker.register_document(
                file_path, 
                'word',
                {'paragraph_count': len(doc.paragraphs), 'table_count': len(doc.tables)}
            )
        
        result = {
            'paragraphs': [],
            'tables': [],
            'headers': [],
            'key_sections': {},
            'raw_text': '',
            'key_metrics': {},
            'document_id': document_id,
            'data_point_ids': {}
        }
        
        # Extract paragraphs with structure
        paragraph_texts = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                para_data = {
                    'id': f'para_{i+1}',
                    'text': paragraph.text.strip(),
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'is_header': self._is_header_style(paragraph)
                }
                
                result['paragraphs'].append(para_data)
                paragraph_texts.append(paragraph.text.strip())
                
                # Identify headers
                if para_data['is_header']:
                    result['headers'].append({
                        'id': para_data['id'],
                        'text': para_data['text'],
                        'level': self._get_header_level(paragraph)
                    })
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = self._extract_table(table, i+1)
            result['tables'].append(table_data)
        
        # Identify key sections
        result['key_sections'] = self._identify_key_sections(result['paragraphs'])
        
        # Create raw text
        result['raw_text'] = '\n\n'.join(paragraph_texts)
        
        # Extract key metrics with source tracking
        result['key_metrics'] = self._extract_key_metrics_with_tracking(
            result['raw_text'], 
            result['paragraphs'], 
            document_id
        )
        
        return result
    
    def _is_header_style(self, paragraph):
        """Determine if paragraph is a header based on style"""
        if not paragraph.style:
            return False
        
        style_name = paragraph.style.name.lower()
        header_indicators = ['heading', 'title', 'subtitle', 'header']
        
        return any(indicator in style_name for indicator in header_indicators)
    
    def _get_header_level(self, paragraph):
        """Get header level from style"""
        if not paragraph.style:
            return 1
        
        style_name = paragraph.style.name.lower()
        
        # Extract number from heading styles like "Heading 1", "Heading 2", etc.
        if 'heading' in style_name:
            try:
                level = int(style_name.split()[-1])
                return level
            except (ValueError, IndexError):
                return 1
        
        return 1
    
    def _extract_table(self, table, table_id):
        """Extract table data with structure"""
        table_data = {
            'id': f'table_{table_id}',
            'rows': [],
            'headers': [],
            'cell_count': 0
        }
        
        for i, row in enumerate(table.rows):
            row_data = []
            for j, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                cell_data = {
                    'text': cell_text,
                    'position': f'{chr(65+j)}{i+1}',  # A1, B1, etc.
                    'is_header': i == 0  # Assume first row is header
                }
                row_data.append(cell_data)
                table_data['cell_count'] += 1
            
            table_data['rows'].append(row_data)
            
            # Store headers from first row
            if i == 0:
                table_data['headers'] = [cell['text'] for cell in row_data]
        
        return table_data
    
    def _identify_key_sections(self, paragraphs):
        """Identify key sections like executive summary, financials, etc."""
        sections = {}
        current_section = None
        
        # Keywords that indicate important sections
        section_keywords = {
            'executive_summary': ['executive summary', 'overview', 'summary'],
            'financial_performance': ['financial', 'revenue', 'profit', 'earnings', 'performance'],
            'growth': ['growth', 'expansion', 'increase', 'trends'],
            'market': ['market', 'industry', 'competitive', 'position'],
            'operations': ['operations', 'business', 'strategy']
        }
        
        for para in paragraphs:
            text_lower = para['text'].lower()
            
            # Check if this paragraph starts a new section
            if para['is_header']:
                for section_type, keywords in section_keywords.items():
                    if any(keyword in text_lower for keyword in keywords):
                        current_section = section_type
                        if current_section not in sections:
                            sections[current_section] = []
                        sections[current_section].append({
                            'paragraph_id': para['id'],
                            'type': 'header',
                            'text': para['text']
                        })
                        break
            elif current_section:
                # Add content to current section
                sections[current_section].append({
                    'paragraph_id': para['id'],
                    'type': 'content',
                    'text': para['text']
                })
        
        return sections
    
    def _extract_key_metrics_with_tracking(self, text: str, paragraphs: list, document_id: Optional[str]) -> Dict[str, Any]:
        """Extract key financial/business metrics from text with paragraph tracking"""
        metrics = {}
        
        # Common patterns for financial metrics
        patterns = {
            'revenue': r'(?:revenue|sales)\s*:?\s*\$?([\d,\.]+[MmBbKk]?)\b',
            'growth': r'(?:growth|increase)\s*:?\s*([\d\.]+%)',
            'profit': r'(?:profit|earnings)\s*:?\s*\$?([\d,\.]+[MmBbKk]?)\b',
            'margin': r'(?:margin)\s*:?\s*([\d\.]+%)',
            'customers': r'(?:customers|clients)\s*:?\s*([\d,]+)\b',
            'employees': r'(?:employees|staff|team)\s*:?\s*([\d,]+)\b',
            'market_share': r'(?:market share)\s*:?\s*([\d\.]+%)',
            'valuation': r'(?:valuation|value)\s*:?\s*\$?([\d,\.]+[MmBbKk]?)\b'
        }
        
        for metric, pattern in patterns.items():
            # Find matches in full text first
            matches = list(re.finditer(pattern, text, re.IGNORECASE))
            if matches:
                match = matches[0]  # Take first match
                value = match.group(1)
                
                # Find which paragraph this match is in
                paragraph_id = self._find_paragraph_for_position(match.start(), text, paragraphs)
                
                # Track data point if source tracker is available
                data_point_id = None
                if self.source_tracker and document_id:
                    confidence = self._calculate_word_confidence(value, metric, match.group(0), paragraph_id)
                    data_point_id = self.source_tracker.track_data_point(
                        value=value,
                        document_id=document_id,
                        location_details={
                            'page_or_sheet': 'Document',
                            'cell_or_section': paragraph_id,
                            'coordinates': {'char_position': match.start()},
                            'extraction_method': 'docx_regex'
                        },
                        confidence=confidence,
                        context=f"Extracted {metric}: {match.group(0)}"
                    )
                
                metrics[metric] = {
                    'value': value,
                    'paragraph': paragraph_id,
                    'data_point_id': data_point_id,
                    'full_match': match.group(0)
                }
        
        return metrics
    
    def _find_paragraph_for_position(self, char_position: int, full_text: str, paragraphs: list) -> str:
        """Find which paragraph a character position belongs to"""
        # Split by double newlines to match how we joined paragraphs
        para_texts = full_text.split('\n\n')
        current_position = 0
        
        for i, para_text in enumerate(para_texts):
            if char_position <= current_position + len(para_text):
                # Find corresponding paragraph ID
                if i < len(paragraphs):
                    return paragraphs[i]['id']
                else:
                    return f'para_{i+1}'
            current_position += len(para_text) + 2  # +2 for the \n\n
        
        # Default to last paragraph if position is beyond content
        if paragraphs:
            return paragraphs[-1]['id']
        else:
            return 'para_1'
    
    def _calculate_word_confidence(self, value: str, metric_type: str, full_match: str, paragraph_id: str) -> float:
        """Calculate confidence score for Word extracted data"""
        confidence = 0.8  # Base confidence for Word extraction (higher than PDF)
        
        # Boost confidence based on value characteristics
        if '$' in full_match:
            confidence += 0.1  # Currency symbol increases confidence
        
        if metric_type in ['revenue', 'profit', 'valuation'] and any(suffix in value.lower() for suffix in ['m', 'b', 'k']):
            confidence += 0.1  # Scale indicators increase confidence
        
        if metric_type in ['growth', 'margin', 'market_share'] and '%' in full_match:
            confidence += 0.1  # Percentage symbol increases confidence
        
        # Reduce confidence for very short values
        if len(value) < 2:
            confidence -= 0.2
        
        # Headers tend to be more reliable
        if 'header' in paragraph_id.lower():
            confidence += 0.05
        
        return max(0.4, min(0.95, confidence))  # Clamp between 0.4 and 0.95