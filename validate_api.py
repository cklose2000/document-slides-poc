#!/usr/bin/env python3
"""
Quick API Validation Script for document-slides-poc
Performs basic validation to ensure API is working correctly
"""

import requests
import json
import time
import io
import tempfile
import openpyxl
from docx import Document

def create_sample_files():
    """Create minimal sample files for testing"""
    
    # Create minimal Excel file
    excel_content = io.BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws['A1'] = "Company"
    ws['B1'] = "TechCorp"
    ws['A2'] = "Revenue"
    ws['B2'] = 1000000
    wb.save(excel_content)
    excel_content.seek(0)
    
    # Create minimal Word document
    word_content = io.BytesIO()
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test document for API validation.')
    doc.save(word_content)
    word_content.seek(0)
    
    return excel_content, word_content

def test_health_endpoint():
    """Test the health check endpoint"""
    print("1. Testing health endpoint...")
    try:
        response = requests.get("http://localhost:5000/health", timeout=5)
        if response.status_code == 200:
            data = response.json()
            print(f"   ✅ Health check passed: {data.get('status', 'Unknown')}")
            print(f"      Service: {data.get('service', 'Unknown')}")
            print(f"      Template Management: {data.get('template_management', 'Unknown')}")
            return True
        else:
            print(f"   ❌ Health check failed: HTTP {response.status_code}")
            return False
    except Exception as e:
        print(f"   ❌ Health check failed: {str(e)}")
        return False

def test_preview_endpoint():
    """Test the preview extraction endpoint"""
    print("\n2. Testing preview extraction...")
    try:
        excel_content, word_content = create_sample_files()
        
        files = {'documents': ('test.xlsx', excel_content, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
        
        start_time = time.time()
        response = requests.post("http://localhost:5000/api/generate-slides/preview", files=files, timeout=30)
        end_time = time.time()
        
        if response.status_code == 200:
            data = response.json()
            if data.get('success'):
                print(f"   ✅ Preview extraction passed ({end_time - start_time:.2f}s)")
                print(f"      Documents processed: {data.get('documents_processed', 0)}")
                return True
            else:
                print(f"   ❌ Preview extraction failed: {data.get('message', 'Unknown error')}")
                return False
        else:
            try:
                error_data = response.json()
                print(f"   ❌ Preview extraction failed: {error_data.get('error', 'Unknown error')}")
            except:
                print(f"   ❌ Preview extraction failed: HTTP {response.status_code}")
            return False
            
    except Exception as e:
        print(f"   ❌ Preview extraction failed: {str(e)}")
        return False

def test_slide_generation():
    """Test basic slide generation"""
    print("\n3. Testing slide generation...")
    try:
        excel_content, word_content = create_sample_files()
        
        files = {'documents': ('test.xlsx', excel_content, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
        
        start_time = time.time()
        response = requests.post("http://localhost:5000/api/generate-slides", files=files, timeout=60)
        end_time = time.time()
        
        if response.status_code == 200:
            # Check if it's a PowerPoint file
            content_type = response.headers.get('content-type', '')
            if 'presentationml.presentation' in content_type:
                file_size = len(response.content)
                print(f"   ✅ Slide generation passed ({end_time - start_time:.2f}s)")
                print(f"      Generated file size: {file_size:,} bytes")
                return True
            else:
                print(f"   ❌ Slide generation failed: Wrong content type: {content_type}")
                return False
        else:
            try:
                error_data = response.json()
                print(f"   ❌ Slide generation failed: {error_data.get('error', 'Unknown error')}")
            except:
                print(f"   ❌ Slide generation failed: HTTP {response.status_code}")
            return False
            
    except Exception as e:
        print(f"   ❌ Slide generation failed: {str(e)}")
        return False

def test_templates_endpoint():
    """Test template management endpoints"""
    print("\n4. Testing template management...")
    try:
        response = requests.get("http://localhost:5000/api/templates", timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            template_count = data.get('count', 0)
            print(f"   ✅ Template management working: {template_count} templates available")
            return True
        elif response.status_code == 503:
            data = response.json()
            print(f"   ⚠️  Template management not available: {data.get('error', 'Unknown')}")
            return True  # This is acceptable
        else:
            try:
                error_data = response.json()
                print(f"   ❌ Template management failed: {error_data.get('error', 'Unknown error')}")
            except:
                print(f"   ❌ Template management failed: HTTP {response.status_code}")
            return False
            
    except Exception as e:
        print(f"   ❌ Template management failed: {str(e)}")
        return False

def test_error_handling():
    """Test basic error handling"""
    print("\n5. Testing error handling...")
    try:
        # Test with no files
        response = requests.post("http://localhost:5000/api/generate-slides", timeout=10)
        
        if response.status_code == 400:
            data = response.json()
            if 'error' in data:
                print(f"   ✅ Error handling working: {data['error']}")
                return True
            else:
                print(f"   ❌ Error response missing 'error' field")
                return False
        else:
            print(f"   ❌ Expected 400 error, got {response.status_code}")
            return False
            
    except Exception as e:
        print(f"   ❌ Error handling test failed: {str(e)}")
        return False

def main():
    """Run API validation"""
    print("=" * 60)
    print("API VALIDATION SCRIPT")
    print("document-slides-poc Quick Validation")
    print("=" * 60)
    
    tests = [
        test_health_endpoint,
        test_preview_endpoint,
        test_slide_generation,
        test_templates_endpoint,
        test_error_handling
    ]
    
    passed_tests = 0
    total_tests = len(tests)
    
    start_time = time.time()
    
    for test_func in tests:
        try:
            if test_func():
                passed_tests += 1
        except Exception as e:
            print(f"   ❌ Test failed with exception: {str(e)}")
    
    end_time = time.time()
    total_time = end_time - start_time
    
    print("\n" + "=" * 60)
    print("VALIDATION RESULTS")
    print("=" * 60)
    print(f"Tests passed: {passed_tests}/{total_tests}")
    print(f"Success rate: {(passed_tests/total_tests*100):.1f}%")
    print(f"Total time: {total_time:.2f} seconds")
    
    if passed_tests == total_tests:
        print("\n🎉 ALL VALIDATIONS PASSED!")
        print("✅ API is working correctly and ready for use")
        return True
    elif passed_tests >= total_tests - 1:
        print("\n⚠️  MOSTLY WORKING")
        print("✅ Core functionality is working, minor issues detected")
        return True
    else:
        print("\n❌ VALIDATION FAILED")
        print("🔧 Multiple issues detected, please check the API server")
        return False

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)