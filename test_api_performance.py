#!/usr/bin/env python3
"""
Performance and Stress Testing Suite for document-slides-poc API
Tests performance, concurrency, memory usage, and stress scenarios
"""

import unittest
import requests
import json
import time
import threading
import queue
import os
import tempfile
import psutil
import gc
from concurrent.futures import ThreadPoolExecutor, as_completed
import statistics
import openpyxl
from docx import Document

class APIPerformanceTestSuite(unittest.TestCase):
    """Performance and stress testing for API endpoints"""
    
    @classmethod
    def setUpClass(cls):
        """Setup performance testing environment"""
        cls.base_url = "http://localhost:5000"
        cls.api_url = f"{cls.base_url}/api"
        
        # Create test data directory
        cls.test_data_dir = tempfile.mkdtemp(prefix="perf_test_")
        
        # Create test files of various sizes
        cls._create_performance_test_files()
        cls.temp_files = []
        
        # Performance tracking
        cls.performance_results = []
    
    @classmethod
    def tearDownClass(cls):
        """Cleanup performance test environment"""
        # Clean up temporary files
        for temp_file in cls.temp_files:
            if os.path.exists(temp_file):
                os.unlink(temp_file)
        
        # Clean up test data directory
        import shutil
        if os.path.exists(cls.test_data_dir):
            shutil.rmtree(cls.test_data_dir)
        
        # Print performance summary
        if cls.performance_results:
            cls._print_performance_summary()
    
    @classmethod
    def _create_performance_test_files(cls):
        """Create test files of various sizes for performance testing"""
        
        # Small Excel file (~50KB)
        cls.small_excel = os.path.join(cls.test_data_dir, "small_data.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Small_Data"
        
        for i in range(50):
            ws[f'A{i+1}'] = f"Item {i}"
            ws[f'B{i+1}'] = i * 100
            ws[f'C{i+1}'] = f"Description {i}"
        
        wb.save(cls.small_excel)
        
        # Medium Excel file (~500KB)
        cls.medium_excel = os.path.join(cls.test_data_dir, "medium_data.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Medium_Data"
        
        # Create multiple sheets with data
        for sheet_num in range(3):
            if sheet_num > 0:
                ws = wb.create_sheet(f"Sheet_{sheet_num}")
            
            for i in range(500):
                ws[f'A{i+1}'] = f"Company {i}"
                ws[f'B{i+1}'] = (i + 1) * 1000
                ws[f'C{i+1}'] = (i + 1) * 50
                ws[f'D{i+1}'] = f"Sector {i % 10}"
                ws[f'E{i+1}'] = f"Location {i % 20}"
        
        wb.save(cls.medium_excel)
        
        # Large Excel file (~2MB)
        cls.large_excel = os.path.join(cls.test_data_dir, "large_data.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Large_Data"
        
        # Create multiple sheets with extensive data
        for sheet_num in range(5):
            if sheet_num > 0:
                ws = wb.create_sheet(f"LargeSheet_{sheet_num}")
            
            for i in range(2000):
                ws[f'A{i+1}'] = f"Record {i}"
                ws[f'B{i+1}'] = (i + 1) * 1000
                ws[f'C{i+1}'] = (i + 1) * 50
                ws[f'D{i+1}'] = (i + 1) * 25
                ws[f'E{i+1}'] = f"Category {i % 50}"
                ws[f'F{i+1}'] = f"Region {i % 25}"
                ws[f'G{i+1}'] = f"Product {i % 100}"
        
        wb.save(cls.large_excel)
        
        # Small Word document
        cls.small_word = os.path.join(cls.test_data_dir, "small_doc.docx")
        doc = Document()
        doc.add_heading('Small Document', 0)
        for i in range(10):
            doc.add_paragraph(f'This is paragraph {i} in a small document for performance testing.')
        doc.save(cls.small_word)
        
        # Medium Word document
        cls.medium_word = os.path.join(cls.test_data_dir, "medium_doc.docx")
        doc = Document()
        doc.add_heading('Medium Document', 0)
        
        for section in range(5):
            doc.add_heading(f'Section {section}', 1)
            for para in range(20):
                doc.add_paragraph(
                    f'This is paragraph {para} in section {section}. '
                    f'It contains detailed information about performance testing '
                    f'and various scenarios that need to be validated. '
                    f'The content is designed to simulate real-world document sizes '
                    f'and complexity for comprehensive testing purposes.'
                )
        
        doc.save(cls.medium_word)
        
        # Large Word document
        cls.large_word = os.path.join(cls.test_data_dir, "large_doc.docx")
        doc = Document()
        doc.add_heading('Large Document', 0)
        
        for section in range(15):
            doc.add_heading(f'Section {section}', 1)
            for subsection in range(3):
                doc.add_heading(f'Subsection {section}.{subsection}', 2)
                for para in range(25):
                    doc.add_paragraph(
                        f'This is paragraph {para} in subsection {section}.{subsection}. '
                        f'Large documents require careful performance optimization '
                        f'to ensure responsive processing times. The system must handle '
                        f'extensive content while maintaining good user experience. '
                        f'This content simulates complex business documents with '
                        f'multiple sections, detailed analysis, and comprehensive coverage '
                        f'of various topics relevant to enterprise document processing.'
                    )
        
        doc.save(cls.large_word)
    
    @classmethod
    def _print_performance_summary(cls):
        """Print performance testing summary"""
        print("\n" + "=" * 60)
        print("PERFORMANCE TESTING SUMMARY")
        print("=" * 60)
        
        if cls.performance_results:
            for category, results in cls.performance_results:
                print(f"\n{category}:")
                if results:
                    avg_time = statistics.mean(results)
                    min_time = min(results)
                    max_time = max(results)
                    median_time = statistics.median(results)
                    
                    print(f"  Average: {avg_time:.2f}s")
                    print(f"  Median:  {median_time:.2f}s")
                    print(f"  Min:     {min_time:.2f}s")
                    print(f"  Max:     {max_time:.2f}s")
                    print(f"  Samples: {len(results)}")
    
    def setUp(self):
        """Setup for each test"""
        self.session = requests.Session()
        self.session.timeout = 120  # Longer timeout for performance tests
        gc.collect()  # Clean up memory before each test
    
    def tearDown(self):
        """Cleanup after each test"""
        self.session.close()
        gc.collect()
    
    def _check_server_running(self):
        """Check if the API server is running"""
        try:
            response = self.session.get(f"{self.base_url}/health", timeout=5)
            return response.status_code == 200
        except:
            return False
    
    def _measure_request_time(self, request_func):
        """Measure the time taken for a request"""
        start_time = time.time()
        result = request_func()
        end_time = time.time()
        return result, end_time - start_time
    
    def _get_memory_usage(self):
        """Get current memory usage"""
        process = psutil.Process()
        return process.memory_info().rss / 1024 / 1024  # MB
    
    def test_01_small_file_performance(self):
        """Test performance with small files"""
        print("\n=== Testing Small File Performance ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        times = []
        
        for i in range(5):  # 5 iterations
            def request():
                with open(self.small_excel, 'rb') as f:
                    files = {'documents': ('small_data.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                    return self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
            
            response, duration = self._measure_request_time(request)
            times.append(duration)
            
            if response.status_code == 200:
                data = response.json()
                self.assertTrue(data['success'])
                print(f"  Iteration {i+1}: {duration:.2f}s - Success")
            else:
                print(f"  Iteration {i+1}: {duration:.2f}s - Error {response.status_code}")
        
        avg_time = statistics.mean(times)
        self.assertLess(avg_time, 5.0)  # Should be under 5 seconds on average
        
        self.performance_results.append(("Small File Performance", times))
        print(f"✓ Small file average time: {avg_time:.2f}s")
    
    def test_02_medium_file_performance(self):
        """Test performance with medium-sized files"""
        print("\n=== Testing Medium File Performance ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        times = []
        
        for i in range(3):  # 3 iterations for medium files
            def request():
                with open(self.medium_excel, 'rb') as f:
                    files = {'documents': ('medium_data.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                    return self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
            
            response, duration = self._measure_request_time(request)
            times.append(duration)
            
            if response.status_code == 200:
                data = response.json()
                self.assertTrue(data['success'])
                print(f"  Iteration {i+1}: {duration:.2f}s - Success")
            else:
                print(f"  Iteration {i+1}: {duration:.2f}s - Error {response.status_code}")
        
        avg_time = statistics.mean(times)
        self.assertLess(avg_time, 15.0)  # Should be under 15 seconds on average
        
        self.performance_results.append(("Medium File Performance", times))
        print(f"✓ Medium file average time: {avg_time:.2f}s")
    
    def test_03_large_file_performance(self):
        """Test performance with large files"""
        print("\n=== Testing Large File Performance ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        times = []
        memory_before = self._get_memory_usage()
        
        for i in range(2):  # 2 iterations for large files
            def request():
                with open(self.large_excel, 'rb') as f:
                    files = {'documents': ('large_data.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                    return self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
            
            response, duration = self._measure_request_time(request)
            times.append(duration)
            
            if response.status_code == 200:
                data = response.json()
                self.assertTrue(data['success'])
                print(f"  Iteration {i+1}: {duration:.2f}s - Success")
            else:
                print(f"  Iteration {i+1}: {duration:.2f}s - Error {response.status_code}")
            
            # Check memory usage
            memory_after = self._get_memory_usage()
            print(f"  Memory usage: {memory_after:.1f}MB (delta: {memory_after - memory_before:.1f}MB)")
        
        avg_time = statistics.mean(times)
        self.assertLess(avg_time, 30.0)  # Should be under 30 seconds on average
        
        self.performance_results.append(("Large File Performance", times))
        print(f"✓ Large file average time: {avg_time:.2f}s")
    
    def test_04_concurrent_requests_performance(self):
        """Test performance under concurrent load"""
        print("\n=== Testing Concurrent Requests Performance ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        def make_concurrent_request(request_id):
            """Make a single concurrent request"""
            start_time = time.time()
            try:
                with open(self.small_excel, 'rb') as f:
                    files = {'documents': (f'concurrent_{request_id}.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                    response = requests.post(f"{self.api_url}/generate-slides/preview", files=files, timeout=60)
                
                end_time = time.time()
                duration = end_time - start_time
                
                return {
                    'request_id': request_id,
                    'status_code': response.status_code,
                    'duration': duration,
                    'success': response.status_code == 200
                }
            except Exception as e:
                end_time = time.time()
                return {
                    'request_id': request_id,
                    'status_code': 'error',
                    'duration': end_time - start_time,
                    'success': False,
                    'error': str(e)
                }
        
        # Test with increasing concurrency levels
        for concurrency in [2, 5, 10]:
            print(f"\n  Testing {concurrency} concurrent requests:")
            
            start_time = time.time()
            
            with ThreadPoolExecutor(max_workers=concurrency) as executor:
                futures = [executor.submit(make_concurrent_request, i) for i in range(concurrency)]
                results = [future.result() for future in as_completed(futures)]
            
            end_time = time.time()
            total_time = end_time - start_time
            
            # Analyze results
            successful = [r for r in results if r['success']]
            failed = [r for r in results if not r['success']]
            
            if successful:
                avg_duration = statistics.mean([r['duration'] for r in successful])
                max_duration = max([r['duration'] for r in successful])
                min_duration = min([r['duration'] for r in successful])
                
                print(f"    Total time: {total_time:.2f}s")
                print(f"    Successful: {len(successful)}/{concurrency}")
                print(f"    Avg request time: {avg_duration:.2f}s")
                print(f"    Min request time: {min_duration:.2f}s")
                print(f"    Max request time: {max_duration:.2f}s")
                
                # Performance assertions
                self.assertGreater(len(successful), concurrency * 0.8)  # At least 80% success rate
                self.assertLess(avg_duration, 15.0)  # Average should be reasonable
                
                self.performance_results.append((f"Concurrent {concurrency} requests", [r['duration'] for r in successful]))
            
            if failed:
                print(f"    Failed requests: {len(failed)}")
                for fail in failed[:3]:  # Show first 3 failures
                    print(f"      Request {fail['request_id']}: {fail.get('error', 'Unknown error')}")
    
    def test_05_mixed_file_types_performance(self):
        """Test performance with mixed file types"""
        print("\n=== Testing Mixed File Types Performance ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        times = []
        
        for i in range(3):
            def request():
                files = [
                    ('documents', ('small_data.xlsx', open(self.small_excel, 'rb'), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')),
                    ('documents', ('small_doc.docx', open(self.small_word, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
                ]
                
                try:
                    return self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
                finally:
                    # Close file handles
                    for _, file_tuple in files:
                        file_tuple[1].close()
            
            response, duration = self._measure_request_time(request)
            times.append(duration)
            
            if response.status_code == 200:
                data = response.json()
                self.assertTrue(data['success'])
                self.assertEqual(data['documents_processed'], 2)
                print(f"  Iteration {i+1}: {duration:.2f}s - Success (2 files)")
            else:
                print(f"  Iteration {i+1}: {duration:.2f}s - Error {response.status_code}")
        
        avg_time = statistics.mean(times)
        self.assertLess(avg_time, 10.0)  # Mixed files should be processed efficiently
        
        self.performance_results.append(("Mixed File Types", times))
        print(f"✓ Mixed file types average time: {avg_time:.2f}s")
    
    def test_06_memory_usage_stability(self):
        """Test memory usage stability over multiple requests"""
        print("\n=== Testing Memory Usage Stability ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        memory_readings = []
        initial_memory = self._get_memory_usage()
        memory_readings.append(initial_memory)
        
        print(f"  Initial memory: {initial_memory:.1f}MB")
        
        # Make multiple requests and track memory
        for i in range(10):
            with open(self.medium_excel, 'rb') as f:
                files = {'documents': (f'memory_test_{i}.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                response = self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
            
            # Allow some time for cleanup
            time.sleep(1)
            gc.collect()
            
            current_memory = self._get_memory_usage()
            memory_readings.append(current_memory)
            
            if i % 3 == 0:  # Print every 3rd iteration
                print(f"  After request {i+1}: {current_memory:.1f}MB")
        
        final_memory = memory_readings[-1]
        memory_growth = final_memory - initial_memory
        max_memory = max(memory_readings)
        
        print(f"  Final memory: {final_memory:.1f}MB")
        print(f"  Memory growth: {memory_growth:.1f}MB")
        print(f"  Peak memory: {max_memory:.1f}MB")
        
        # Memory growth should be reasonable (less than 100MB for this test)
        self.assertLess(memory_growth, 100.0)
        print("✓ Memory usage appears stable")
    
    def test_07_response_time_consistency(self):
        """Test response time consistency for identical requests"""
        print("\n=== Testing Response Time Consistency ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        times = []
        
        # Make 10 identical requests
        for i in range(10):
            def request():
                with open(self.small_excel, 'rb') as f:
                    files = {'documents': ('consistency_test.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                    return self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
            
            response, duration = self._measure_request_time(request)
            times.append(duration)
            
            if response.status_code == 200:
                print(f"  Request {i+1}: {duration:.2f}s")
            else:
                print(f"  Request {i+1}: {duration:.2f}s - Error {response.status_code}")
        
        # Analyze consistency
        if len(times) >= 5:
            avg_time = statistics.mean(times)
            std_dev = statistics.stdev(times)
            coefficient_of_variation = std_dev / avg_time if avg_time > 0 else 0
            
            print(f"  Average time: {avg_time:.2f}s")
            print(f"  Standard deviation: {std_dev:.2f}s")
            print(f"  Coefficient of variation: {coefficient_of_variation:.2f}")
            
            # Response times should be reasonably consistent (CV < 0.5)
            self.assertLess(coefficient_of_variation, 0.5)
            
            self.performance_results.append(("Response Time Consistency", times))
            print("✓ Response times are consistent")
    
    def test_08_error_handling_performance(self):
        """Test performance of error handling"""
        print("\n=== Testing Error Handling Performance ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        error_times = []
        
        # Test various error scenarios
        error_scenarios = [
            ("No files", lambda: self.session.post(f"{self.api_url}/generate-slides")),
            ("Invalid file", lambda: self._test_invalid_file()),
            ("Nonexistent template", lambda: self._test_nonexistent_template())
        ]
        
        for scenario_name, scenario_func in error_scenarios:
            start_time = time.time()
            try:
                response = scenario_func()
                end_time = time.time()
                duration = end_time - start_time
                error_times.append(duration)
                
                # Should return error quickly
                self.assertLess(duration, 5.0)
                print(f"  {scenario_name}: {duration:.2f}s - Status {response.status_code}")
                
            except Exception as e:
                end_time = time.time()
                duration = end_time - start_time
                print(f"  {scenario_name}: {duration:.2f}s - Exception: {str(e)}")
        
        if error_times:
            avg_error_time = statistics.mean(error_times)
            print(f"✓ Average error handling time: {avg_error_time:.2f}s")
            self.performance_results.append(("Error Handling", error_times))
    
    def _test_invalid_file(self):
        """Helper method to test invalid file upload"""
        invalid_content = b"This is not a valid file"
        files = {'documents': ('invalid.txt', invalid_content, 'text/plain')}
        return self.session.post(f"{self.api_url}/generate-slides", files=files)
    
    def _test_nonexistent_template(self):
        """Helper method to test nonexistent template"""
        with open(self.small_excel, 'rb') as f:
            files = {'documents': ('test.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
            data = {'template_id': 'nonexistent_template'}
            return self.session.post(f"{self.api_url}/generate-slides", files=files, data=data)
    
    def test_09_stress_test_rapid_requests(self):
        """Stress test with rapid sequential requests"""
        print("\n=== Stress Testing - Rapid Sequential Requests ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        times = []
        errors = 0
        
        # Make 20 rapid requests
        for i in range(20):
            start_time = time.time()
            
            try:
                with open(self.small_excel, 'rb') as f:
                    files = {'documents': (f'stress_{i}.xlsx', f, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
                    response = self.session.post(f"{self.api_url}/generate-slides/preview", files=files)
                
                end_time = time.time()
                duration = end_time - start_time
                times.append(duration)
                
                if response.status_code != 200:
                    errors += 1
                
                if i % 5 == 0:  # Print every 5th request
                    print(f"  Request {i+1}: {duration:.2f}s - Status {response.status_code}")
                
            except Exception as e:
                errors += 1
                print(f"  Request {i+1}: Error - {str(e)}")
        
        success_rate = ((20 - errors) / 20) * 100
        avg_time = statistics.mean(times) if times else 0
        
        print(f"  Success rate: {success_rate:.1f}%")
        print(f"  Average response time: {avg_time:.2f}s")
        print(f"  Errors: {errors}/20")
        
        # Should handle rapid requests reasonably well
        self.assertGreater(success_rate, 70.0)  # At least 70% success rate
        
        if times:
            self.performance_results.append(("Stress Test - Rapid Requests", times))
        
        print("✓ Stress test completed")
    
    def test_10_health_check_performance(self):
        """Test health check endpoint performance"""
        print("\n=== Testing Health Check Performance ===")
        
        if not self._check_server_running():
            self.skipTest("API server not running")
        
        times = []
        
        # Health check should be very fast
        for i in range(10):
            start_time = time.time()
            response = self.session.get(f"{self.base_url}/health")
            end_time = time.time()
            
            duration = end_time - start_time
            times.append(duration)
            
            self.assertEqual(response.status_code, 200)
            
            if i < 3:  # Print first 3
                print(f"  Health check {i+1}: {duration:.3f}s")
        
        avg_time = statistics.mean(times)
        max_time = max(times)
        
        print(f"  Average: {avg_time:.3f}s")
        print(f"  Maximum: {max_time:.3f}s")
        
        # Health check should be very fast
        self.assertLess(avg_time, 0.1)  # Should be under 100ms on average
        self.assertLess(max_time, 0.5)  # Should never take more than 500ms
        
        self.performance_results.append(("Health Check", times))
        print("✓ Health check performance is good")

def run_performance_tests():
    """Run the performance test suite"""
    print("=" * 60)
    print("PERFORMANCE TESTING SUITE")
    print("Testing API performance, concurrency, and stress scenarios")
    print("=" * 60)
    
    # Create test suite
    loader = unittest.TestLoader()
    suite = loader.loadTestsFromTestCase(APIPerformanceTestSuite)
    
    # Run tests with detailed output
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    
    # Print summary
    print("\n" + "=" * 60)
    print("PERFORMANCE TEST RESULTS")
    print("=" * 60)
    print(f"Tests run: {result.testsRun}")
    print(f"Failures: {len(result.failures)}")
    print(f"Errors: {len(result.errors)}")
    print(f"Skipped: {len(result.skipped)}")
    
    if result.failures:
        print("\nPERFORMANCE FAILURES:")
        for test, traceback in result.failures:
            print(f"- {test}")
    
    if result.errors:
        print("\nPERFORMANCE ERRORS:")
        for test, traceback in result.errors:
            print(f"- {test}")
    
    success_rate = ((result.testsRun - len(result.failures) - len(result.errors)) / result.testsRun * 100) if result.testsRun > 0 else 0
    print(f"\nSuccess Rate: {success_rate:.1f}%")
    
    return result.wasSuccessful()

if __name__ == "__main__":
    success = run_performance_tests()
    exit(0 if success else 1)